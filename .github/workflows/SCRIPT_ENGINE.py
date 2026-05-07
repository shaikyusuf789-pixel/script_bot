# -*- coding: utf-8 -*-
# ============================================================
# SCRIPT_ENGINE.py — SKY Academy Video Script Generator v4.0
# Three Input Modes: Topic | Multi-Transcript Merge | Book PDF
# Video Types: General (Strategy/Motivation) | Subjective (Deep Teaching)
# v4.0: Telugu Lipi Output · Script-Only (No Image Prompts) · 150-180w Chunks
#        AI Thumbnail Generator (DALL-E 3) · Standalone Handout Mode
# SKY Academy Internal Tool
# ============================================================

import streamlit as st
import json
import re
import math
import io
import base64
import random

st.set_page_config(
    page_title="🎬 SKY Academy – Script Engine",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
    .sky-header {
        background: linear-gradient(135deg, #020024 0%, #090979 40%, #00d4ff 100%);
        padding: 30px 20px 22px; border-radius: 18px; text-align: center;
        margin-bottom: 24px; border: 2px solid rgba(255,215,0,0.55);
        box-shadow: 0 8px 48px rgba(0,180,255,0.28), inset 0 0 80px rgba(0,0,0,0.25);
    }
    .sky-logo { font-size:3.6rem; font-weight:900; letter-spacing:8px; margin:0; line-height:1; }
    .sky-s { color:#FF6B6B; text-shadow:0 0 28px #FF6B6B,0 0 56px rgba(255,107,107,0.45); }
    .sky-k { color:#FFE66D; text-shadow:0 0 28px #FFE66D,0 0 56px rgba(255,230,109,0.45); }
    .sky-y { color:#4ECDC4; text-shadow:0 0 28px #4ECDC4,0 0 56px rgba(78,205,196,0.45); }
    .sky-acad { color:rgba(255,255,255,0.92); font-size:0.88rem; font-weight:800;
        letter-spacing:3px; margin:4px 0 0; text-transform:uppercase; }
    .sky-tagline { color:#FFE66D; font-size:1.55rem; font-weight:800; margin:10px 0 3px; }
    .sky-sub { color:rgba(185,225,255,0.88); font-size:0.82rem; margin:4px 0 12px; }
    .sky-pills { display:flex; justify-content:center; gap:8px; flex-wrap:wrap; }
    .sky-pill { background:rgba(255,255,255,0.13); color:#fff;
        border:1px solid rgba(255,255,255,0.28); border-radius:20px;
        padding:3px 14px; font-size:0.74rem; font-weight:600; }
    .mode-topic { background:linear-gradient(to right,#eff6ff,#dbeafe);
        border-left:5px solid #3b82f6; padding:12px 16px;
        border-radius:0 10px 10px 0; font-size:0.84rem; color:#1e40af; margin-bottom:14px; }
    .mode-transcript { background:linear-gradient(to right,#fffbeb,#fef3c7);
        border-left:5px solid #f59e0b; padding:12px 16px;
        border-radius:0 10px 10px 0; font-size:0.84rem; color:#78350f; margin-bottom:14px; }
    .mode-pdf { background:linear-gradient(to right,#ecfdf5,#d1fae5);
        border-left:5px solid #10b981; padding:12px 16px;
        border-radius:0 10px 10px 0; font-size:0.84rem; color:#065f46; margin-bottom:14px; }
    .fcard-ok { background:#f0fdf4; border:1px solid #86efac; border-radius:10px;
        padding:8px 14px; margin:4px 0; font-size:0.82rem; color:#14532d; }
    .fcard-err { background:#fff1f2; border:1px solid #fca5a5; border-radius:10px;
        padding:8px 14px; margin:4px 0; font-size:0.82rem; color:#991b1b; }
    .merge-box { background:linear-gradient(to right,#fefce8,#fef9c3);
        border:1px solid #fde047; border-radius:10px; padding:10px 14px;
        font-size:0.83rem; color:#713f12; margin:8px 0; }
    .stat-pill { display:inline-block; background:linear-gradient(135deg,#e8e0ff,#dbeafe);
        color:#302b63; border-radius:20px; padding:3px 12px;
        font-size:0.8rem; margin:2px; font-weight:600; }
    .badge-topic { background:linear-gradient(135deg,#3b82f6,#1d4ed8); color:#fff;
        border-radius:8px; padding:4px 12px; font-size:.8rem; font-weight:700; }
    .badge-transcript { background:linear-gradient(135deg,#f59e0b,#d97706); color:#fff;
        border-radius:8px; padding:4px 12px; font-size:.8rem; font-weight:700; }
    .badge-pdf { background:linear-gradient(135deg,#10b981,#059669); color:#fff;
        border-radius:8px; padding:4px 12px; font-size:.8rem; font-weight:700; }
    .word-info { background:linear-gradient(to right,#f0edff,#e8e0ff);
        border-left:5px solid #302b63; padding:10px 14px;
        border-radius:0 10px 10px 0; font-size:0.85rem; color:#302b63; margin:4px 0 10px 0; }
    .stream-progress { background:linear-gradient(to right,#eff6ff,#dbeafe);
        border-left:4px solid #3b82f6; padding:10px 16px;
        border-radius:0 10px 10px 0; font-size:0.84rem; color:#1e40af; margin:8px 0; }
    .live-stream-box {
        background: linear-gradient(135deg, #0f0c29 0%, #302b63 60%, #24243e 100%);
        border: 2px solid rgba(167,139,250,0.5);
        border-radius: 14px; padding: 18px 20px; margin: 10px 0 16px 0;
        box-shadow: 0 6px 32px rgba(120,0,255,0.25);
    }
    .live-stream-title {
        color: #e9d5ff; font-size: 0.85rem; font-weight: 800;
        letter-spacing: 3px; text-transform: uppercase; margin-bottom: 10px;
    }
    .live-stream-hint {
        color: #a78bfa; font-size: 0.75rem; margin-bottom: 10px; font-style: italic;
    }
    .empty-preview { background:linear-gradient(135deg,#f8faff,#f0f2f6);
        border:2px dashed #c7d2fe; border-radius:16px; padding:60px 20px;
        text-align:center; color:#6366f1; }
    .key-ok   { color:#0da271; font-size:11px; margin-top:-6px; }
    .key-warn { color:#f59e0b; font-size:11px; margin-top:-6px; }
    .vtype-general { background:linear-gradient(to right,#fff7ed,#ffedd5);
        border-left:5px solid #f97316; padding:10px 14px;
        border-radius:0 10px 10px 0; font-size:0.82rem; color:#7c2d12; margin:8px 0 12px 0; }
    .vtype-subjective { background:linear-gradient(to right,#f0fdf4,#dcfce7);
        border-left:5px solid #22c55e; padding:10px 14px;
        border-radius:0 10px 10px 0; font-size:0.82rem; color:#14532d; margin:8px 0 12px 0; }
    .seo-box { background:linear-gradient(to right,#fdf4ff,#fae8ff);
        border-left:5px solid #a855f7; padding:12px 16px;
        border-radius:0 10px 10px 0; font-size:0.83rem; color:#581c87; margin:10px 0; }
    .regen-hint { background:linear-gradient(to right,#eff6ff,#e0f2fe);
        border:1px dashed #7dd3fc; border-radius:8px; padding:8px 12px;
        font-size:0.78rem; color:#075985; margin-bottom:6px; }
    .handout-block {
        background: linear-gradient(135deg, #1e0040 0%, #2d0060 50%, #4a0080 100%);
        border: 2px solid rgba(167,139,250,0.45);
        border-radius: 18px; padding: 24px 28px; margin-top: 28px;
        box-shadow: 0 8px 40px rgba(120,0,255,0.2);
    }
    .handout-title {
        color: #e9d5ff; font-size: 1.3rem; font-weight: 800;
        letter-spacing: 2px; margin-bottom: 16px; text-align: center;
    }
    .handout-hint {
        background: rgba(167,139,250,0.15); border: 1px solid rgba(167,139,250,0.3);
        border-radius: 8px; padding: 10px 14px; font-size: 0.8rem;
        color: #c4b5fd; margin-bottom: 14px;
    }
    .parse-warn { background:linear-gradient(to right,#fff7ed,#ffedd5);
        border-left:5px solid #f97316; padding:10px 14px;
        border-radius:0 10px 10px 0; font-size:0.82rem; color:#7c2d12; margin:8px 0; }
    .thumb-block {
        background: linear-gradient(135deg, #1a0800 0%, #2d1500 60%, #4a2500 100%);
        border: 2px solid rgba(251,191,36,0.55);
        border-radius: 18px; padding: 24px 28px; margin-top: 28px;
        box-shadow: 0 8px 40px rgba(180,100,0,0.22);
    }
    .thumb-title {
        color: #fbbf24; font-size: 1.3rem; font-weight: 800;
        letter-spacing: 2px; margin-bottom: 16px; text-align: center;
    }
    .thumb-hint {
        background: rgba(251,191,36,0.12); border: 1px solid rgba(251,191,36,0.3);
        border-radius: 8px; padding: 10px 14px; font-size: 0.8rem;
        color: #fde68a; margin-bottom: 14px;
    }
    .thumb-sugg-box {
        background: rgba(251,191,36,0.08); border: 1px solid rgba(251,191,36,0.25);
        border-radius: 8px; padding: 8px 12px; font-size: 0.78rem;
        color: #fef3c7; margin: 4px 0;
    }
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg,#302b63,#0f0c29) !important;
        color:white !important; border:none !important;
        font-weight:600 !important; border-radius:8px !important;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="sky-header">
    <div class="sky-logo">
        <span class="sky-s">S</span><span class="sky-k">K</span><span class="sky-y">Y</span>
    </div>
    <div class="sky-acad">ACADEMY</div>
    <div class="sky-tagline">SCRIPT ENGINE v4.0</div>
    <div class="sky-sub">Telugu Lipi Output &nbsp;·&nbsp; Script-Only &nbsp;·&nbsp; 150-180w Chunks &nbsp;·&nbsp; Thumbnail Generator &nbsp;·&nbsp; Standalone Handout</div>
    <div class="sky-pills">
        <span class="sky-pill">📌 Topic → Original</span>
        <span class="sky-pill">📝 Multi-Transcript → Merge</span>
        <span class="sky-pill">📚 Book PDF → SKY Voice</span>
        <span class="sky-pill">🎯 General · 📖 Subjective</span>
        <span class="sky-pill">🔄 Per-Segment Regen</span>
        <span class="sky-pill">🎬 SEO Pack</span>
        <span class="sky-pill">🖼️ AI Thumbnail</span>
        <span class="sky-pill">📄 AI Handout</span>
    </div>
</div>
""", unsafe_allow_html=True)

# ============================================================
# CONSTANTS
# ============================================================
PAGEGRID_BASE_URL     = "https://api.pagegrid.in"
HARDCODED_CLAUDE_KEY  = "sk-pgrid-bc7765111cae9b0aa9cd62125ff078523667bd377533de3d0396ad2bdda5727b"
WORDS_PER_SEGMENT     = 165
SHEET_ID              = "1dNHDgkX6vhdhZSi5SavBgNihWe04zayRQwyMcCwNlOI"

_HARDCODED_CREDS = {
    "type": "service_account",
    "project_id": "gen-lang-client-0268678328",
    "private_key_id": "06350e4cf0251abf195e1ad49cc5eaf5827ba5ca",
    "private_key": (
        "-----BEGIN PRIVATE KEY-----\n"
        "MIIEvgIBADANBgkqhkiG9w0BAQEFAASCBKgwggSkAgEAAoIBAQCkxNptNHDP8qKi\n"
        "JSqkGbx3SvgY1gnlUoKGzJoldMYvVb6Yw4j7fg/+G5h9P0ze2GtFFD97A7yaUsk0\n"
        "VaH4ODeB5cf414QYP4WU3FoyvgcefkAlGANcZEOt+Enb3K+hhd1bJjfqFXiBO9QT\n"
        "AmsORofIwKJUhdt4+it/4XlC+ADJFd/JKvWLvGUK0p6Aqaz9cgftXoBybF7nTelE\n"
        "j1BofII0hKZ3jrV767SJ7k6qXNVZRydMQY4cqua/8yuVpkM6ALiJZF8/n9Q317y3\n"
        "/DOKgHV4SMno2jf/8MvnIwQExJH5uv0TWo+iFxQ794J0ytpz+Pya83gYFX2eIo6Z\n"
        "gnBnpUnXAgMBAAECggEACN/5i+xJL0o4bFdoJpKkTiChoGTW/50kHrKikuXpTt9l\n"
        "dsEBfdpabit6WTSxpUcu7/eZO70FyaIv6Du8j6wngT2pOcQR/2Rcg5oi2ZzWsVPH\n"
        "jLfwZmeYJaS8BbWrWB3nwGMcm+UwKnXYhHWa4pf19GA73iWfnrKK6UZxy6OkFzCa\n"
        "QyCHEjjNTG/8oCSz5caUCef7IfmaFqu/sg7A2cTekYyltDdYEmWUOx+9nCzvD/KQ\n"
        "wPqkn6jxWuZ2+Ec9+gUuISHpc9G248Rgi97LCufh0JI4NV+zNMG0bEtwz0IpL9sA\n"
        "Q76OoxcdIFoNJSSrUx/F+Ut7hnYJ/1zbH4oPa1CjIQKBgQDbPLgiYeR96nZ24DjI\n"
        "mDVNyhZTauIBbSZR7/u5jI3P5VLLMbfW93hkBQ/3P0IIIJDJIzoHOOnTnNn3KhJE\n"
        "seLjwg9djJKtAtjIWKOXiiOUz/gHnuLnfnmwyKtQnbASfVEjo5AUifVMF90/VkSE\n"
        "elHw6iPA/tE06f9Bm0V9QzaIGQKBgQDAZfR6fPZAYfiUG+u9fUeeZkpGtIXdPWTF\n"
        "wg0RrBa0Kh7o/T35OVgAQnDtG7odE3akrxenytacyzc87IFpC+BSgf+t4mVJnqPY\n"
        "XH5BFORLhVuUoUnNNYq/BF6sY9GZ4e35b+Baxn25GGntT/9pkihYDPPql8Z/t9jb\n"
        "1NCttldfbwKBgEMWvqZO3JQnpp7UGKxR36XxXImkYIrdMufKD3cFavQekgp6KW7Q\n"
        "BfhdkDgyFGvWQ1g5vm0tXmiSTCUVq8d3xB28aeVPuibVgy8z6MPb0u2cAqOaXIdI\n"
        "rcaKcdpWluXhkW3dhJ60ZOsnNl5GcOs1X1Pg4pYRpEWUAbe64zXk1pApAoGBALx2\n"
        "4r3djMbScVJ76zeJ8c7a+mU6TmrCyeThyjWGchL3s6Gc98ka//X5H29UGsKCn1SA\n"
        "Y1as3f9nHOvj7Hw+8vU/fHoTbA5qhKrbJ52O3naP4n68Y3PNv+SPXkHV4aqwYpFV\n"
        "otqo1tyqapDZLSN31WczAPfKxtmy+I2WcPfIxtunAoGBAIVXg/TVXY7KZL/ZLVHx\n"
        "AP2rq/12zXeuUiC13MpZ6RtbiWfNwiPd2HjCNpwjyAxuLjPTTqYUVFZKaKhgvB8/\n"
        "IUQmz1dnM8mdB27jLzjchPSeJFH6yZQKklJqbtopwN557j3vyFIrcxNQJcgpj9UD\n"
        "VyLML2CO4hHr3eoKQ+BNmZj6\n"
        "-----END PRIVATE KEY-----\n"
    ),
    "client_email": "forscripting@gen-lang-client-0268678328.iam.gserviceaccount.com",
    "client_id": "110013612770034478540",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": (
        "https://www.googleapis.com/robot/v1/metadata/x509/"
        "forscripting%40gen-lang-client-0268678328.iam.gserviceaccount.com"
    ),
    "universe_domain": "googleapis.com",
}

MODEL_OPTIONS = {
    "☁️  Claude  (via PageGrid)": ["claude-opus-4-6","claude-sonnet-4-6","claude-haiku-4-5"],
    "🟢  OpenAI  (GPT)":          ["o3","o1","gpt-4.5-preview","gpt-4o"],
    "🔵  Google  (Gemini)":       ["gemini-2.5-pro","gemini-2.0-pro-exp","gemini-2.0-flash","gemini-1.5-pro"],
}

# ============================================================
# SKY ACADEMY STYLE DNA
# ============================================================
_SKY_DNA = """
================================================================
MANDATORY RULE 0 -- TELUGU LIPI (UNICODE SCRIPT) -- HIGHEST PRIORITY
================================================================

THE telugu_text FIELD MUST USE TELUGU UNICODE CHARACTERS FOR ALL TELUGU WORDS.
NEVER USE ROMAN TRANSLITERATION (English letters to spell Telugu words).

WRONG OUTPUT (transliteration -- will be REJECTED):
  "so friends, eeroju manam Preamble gurinchi chuddam ookenaa"
  "idi chala important -- note chesukokandi"
  "meeru idi miss avvadam cheyyaddu"

CORRECT OUTPUT (Telugu Unicode lipi -- REQUIRED):
  "so friends, ఈరోజు మనం Preamble గురించి చూద్దాం ఓకేనా"
  "ఇది చాలా important -- note చేసుకోండి"
  "మీరు ఇది miss అవ్వడం చేయద్దు"

RULE BREAKDOWN:
  Telugu vocabulary & grammar words → Telugu Unicode script ALWAYS
  English technical terms (Article, UPSC, Amendment, SSC CGL) → English letters
  Borrowed English words in everyday Telugu speech → Telugu script:
    "exam"      → "ఎగ్జామ్"
    "important" → "ఇంపార్టెంట్"
    "clear"     → "క్లియర్"
    "strategy"  → "స్ట్రాటెజీ"
    "power"     → "పవర్"
    "security"  → "సెక్యూరిటీ"
    "note it down" → stays as "note it down" (common English classroom phrase)
    "lets go"   → stays as "lets go"
    "friends"   → stays as "friends" (common address in AP Telugu classrooms)
    "Got it?"   → stays as "Got it?"

NOTE: The content examples throughout this DNA use Roman letters to ILLUSTRATE
CONTENT PATTERNS only. Your actual telugu_text output MUST use Telugu Unicode.

================================================================
CORE PHILOSOPHY -- READ THIS FIRST
================================================================

You are a REAL tutor from Andhra Pradesh standing in front of students.
You are NOT translating a textbook.
You are THINKING in Telugu first, then speaking.

ABSOLUTE RULE -- ELEVENLABS TTS COMPATIBILITY:
  The "telugu_text" field MUST contain ZERO emoji characters.
  ElevenLabs will ERROR OUT or read emoji names aloud if emojis are present.

THE MOST IMPORTANT RULE:
  MEANING COMES FIRST. STYLE COMES SECOND.
  Every sentence must mean something on its own.
  Never sacrifice the explanation for a connector phrase.
  Never leave a gap in context between two sentences.

================================================================
NATURAL AP TUTOR VOICE
================================================================

BAD -- bookish, hollow, meaningless:
  "India ni Sovereign, Socialist, Secular, Democratic Republic ga chesyalanukunnaaru"
  "ee concept yokka importance artham chesukovalante manam history loki vellali"

GOOD -- natural, contextual, meaningful AP tutor style:
  [In Telugu lipi]: Set up WHY, then say WHAT, then give exam angle.
  Use real years, real examples, real humor.
  Build suspense only when there is something to reveal.

================================================================
DELIVERY CUES
================================================================
[Energetic]  [Serious]  [Whisper/Secret Tip]  [High Pitch]
[Laughing]   [Deep Pause]  [Assertive]  [Calm, Instructional]
[WARM, FRIENDLY -- WELCOME]

CONNECTOR PHRASES (write in Telugu script in actual output):
"ఓకేనా"  "ఓకే right"  "అవునా కాదా"  "చాలా important"
"తెలుసు కదా"  "మీకు తెలుసు కదా"  "Clear గా అర్థమైందా?"
"lets go"  "note it down"  "Got it?"

LANGUAGE STYLE:
- Telugu Unicode script for all Telugu + borrowed English words
- English for: exam codes (SSC CGL, UPSC), article numbers, amendment names, technical terms
- Direct address in Telugu script: "మీరు", "మీకు", "చూడండి", "అర్థమైందా"
- "friends" stays in English (universal classroom address)
- Light humor only when it fits -- never forced

================================================================
NUMBERS -- ALWAYS WRITE IN WORDS IN telugu_text -- MANDATORY
================================================================

ElevenLabs TTS CANNOT reliably pronounce numerals written as digits.
ALWAYS write ALL numbers as ENGLISH WORDS.

BAD: "1947 లో Independence వచ్చింది"
GOOD: "nineteen forty-seven లో Independence వచ్చింది"

BAD: "Article 21 fundamental right"
GOOD: "Article twenty-one fundamental right"

BAD: "42nd Amendment 1976 లో వచ్చింది"
GOOD: "forty-second Amendment nineteen seventy-six లో వచ్చింది"

NUMBER WORD RULES:
- Years: "nineteen forty-seven", "two thousand twenty-six"
- Article numbers: "Article twenty-one", "Article three-five-six"
- Amendment numbers: "forty-second Amendment", "seventy-third Amendment"
- Percentages: "thirty percent"
- Money: "five thousand rupees", "one lakh rupees"
- Ordinals: "first", "second", "forty-second" (NEVER "42nd")
- Schedules: "sixth Schedule", "tenth Schedule"

================================================================
TELUGU PRONUNCIATION -- DIRGHAM (LONG VOWELS) -- MANDATORY
================================================================

ALWAYS verify correct Telugu spelling.
Long-a = aa sound (ా), Long-i = ii/ee sound (ీ), Long-u = uu/oo sound (ూ)

CRITICAL words:
  కూచిపూడి (KUU-chi-PUU-di) -- both syllables long
  పూజ (PUU-ja) -- long u
  If unsure: write the word in English rather than guess wrong Telugu spelling

================================================================
LIST AND SERIES -- HUMANIZED FLOW
================================================================

NEVER dump items as flat robotic sequence.
ALWAYS use connecting language between items.

BAD: "Pillar one Visual Arts. Pillar two Architecture. Pillar three Literature."

GOOD: "First pillar లో ఏముందంటే -- Visual Arts, Architecture, Sculpture --
eyes తో enjoy చేసే అన్నీ ఇక్కడే cover అవుతాయి ఓకేనా.
Next pillar కి వస్తే -- second one complete గా Literature కి reserved --
books, poems, writing forms -- అన్నీ ఇక్కడే!"

TRANSITION PHRASES between list items:
"next కి వస్తే --" / "ఒక step ముందుకు --" / "ఇంకా interesting గా --"
"ఈ part లో highlight ఏమిటంటే --" / "final గా --" / "చివరగా --"

================================================================
MEMORY HINTS -- RULES
================================================================

ABSOLUTE BAN: NO abbreviation mnemonics, NO first-letter tricks, NO acronyms.
ONLY use WORD-ASSOCIATION and INTERLINKING.

NEVER say "Memory Hint" or "Memory Trick" in the script.
Introduce naturally:
  "దీన్ని simple గా --"
  "ఎలా గుర్తు పెట్టుకోవాలంటే --"
  "best trick ఏమిటంటే --"
  "ఒక simple connection చూడండి --"

EXAMPLES (content pattern; write actual output in Telugu lipi):
  "forty-two degrees fever → forty-second Amendment, Emergency era"
  "Article twenty-one, twenty-first birthday = life's biggest day = Right to LIFE"
  "tenth Schedule = Anti-Defection = padi = fell = politician who betrays party FALLS"
  "Rajya Sabha never fully dissolves → like a marriage, no complete divorce ever"

================================================================
STRICT RULES FOR ALL MODES
================================================================
1. NO bookish headings inside voiceover (SECTION 1, CHAPTER, PART N, etc.)
2. Script flows as ONE continuous natural conversation
3. NEVER mention competitor channels, instructors, books, apps by name
4. ONLY SKY Academy -- weave "SKY Academy లో" naturally where it fits
5. LAST SEGMENT must close with SKY Academy CTA
6. ZERO emojis in telugu_text -- hard technical requirement
7. ALL numbers written as words -- hard technical requirement
8. ALL Telugu words in Telugu Unicode script -- hard technical requirement
"""

TELUGU_TTS_MASTER_PROMPT = """
================================================================
Telugu Voice Output -- Critical Rules
================================================================

All Telugu content in telugu_text must use Telugu Unicode script characters.
Never Roman transliteration for Telugu words.

Consonant stress must be clear:
cha (చ), chha (ఛ), ta (ట), tha (ఠ), da (ద), dha (ధ)
Double consonants (ottulu) must be stressed.
ksha (క్ష) and gya (జ్ఞ) must be pronounced clearly.

Voice Style:
- Speak like a teacher in a classroom -- slow, clear
- One emotion tag per one or two lines: [Energetic] [Serious] [Calm, Instructional]
- Natural pauses, no robotic style

Punctuation and Flow:
- Use -- (double dash) for natural pauses
- Questions must be answered immediately
- Every line must be meaningful
"""

_OUTPUT_FORMAT = """
================================================================
HOOK + WELCOME -- MANDATORY FOR SEGMENT 1 ONLY
================================================================
Segment 1 MUST ALWAYS open with:

ONE POWERFUL HOOK LINE (in Telugu Unicode script -- stops the student from scrolling).
WELCOME LINE immediately follows:
   [WARM, FRIENDLY -- WELCOME] Hello Everyone, Welcome to sky academy, where you not only learn the subject but memorise it for ever. [Energetic] Lets start!
Then flow DIRECTLY into content.

================================================================
CLOSING CTA -- MANDATORY FOR LAST SEGMENT
================================================================
Last segment must end with (write Telugu portions in Telugu script):

[Energetic] so friends -- ఈరోజు మనం [TOPIC] గురించి చాలా deep గా చూసుకున్నాం ఓకేనా.
ఈ video యొక్క classroom notes -- print చేసుకోడానికి ready గా ఉన్న complete study material --
SKY Academy Telegram channel లో share చేస్తాను, video చూసిన తర్వాత download చేసుకుని
revision కి use చేయండి -- absolutely free!
మీరు ఇంకా ఏ topic కావాలో, ఏ subject మీద video కావాలో -- comment section లో చెప్పండి.
నేను personally ప్రతి comment చదువుతాను and reply ఇస్తాను -- ఇది నా word మీకు!
Any doubts ఉన్నాయా? Comment below -- I will answer each one personally!

================================================================
OUTPUT FORMAT -- RETURN VALID JSON ARRAY ONLY
================================================================

Return ONLY a valid JSON array. No preamble, no markdown fences, no explanation text.
[
  {
    "seg": 1,
    "title": "3-5 word English heading for this segment (for study notes organization)",
    "telugu_text": "full voiceover in TELUGU UNICODE SCRIPT -- ALL Telugu words in Telugu lipi -- NO Roman transliteration -- NO EMOJIS -- ALL NUMBERS AS ENGLISH WORDS"
  },
  ...
]
- Generate exactly {NUM_SEGS} segments
- Each segment MUST be 150-180 words (STRICTLY -- do not go below 150 or above 180)
- ZERO emojis in telugu_text
- ALL Telugu words in Telugu Unicode script (ఇలాంటివి) -- NEVER Roman (like this: "idi important")
- ALL numbers written as English words (nineteen forty-seven, not 1947)
- Ensure the JSON array is complete and properly closed with ]
"""

_DNA_GENERAL_VIDEO = """
================================================================
VIDEO TYPE: GENERAL -- STRATEGY / GUIDANCE / MOTIVATION
================================================================
High motivation energy. Think passionate senior talking to juniors.
"SKY Academy మీతో ఉంది -- మీరు ఒంటరిగా లేరు."
Max three or four strategy memory hints. Community building. Telegram CTA.
"""

_DNA_SUBJECTIVE_VIDEO = """
================================================================
VIDEO TYPE: SUBJECTIVE -- DEEP SUBJECT TEACHING
================================================================
MINIMUM ONE memory hint per major concept.
After each concept: mention PYQ angle naturally.
Last segment: SKY Academy app CTA + Telegram study notes CTA.
"""

_SYSTEM_TOPIC = (
    "You are an expert Telugu video script writer for SKY Academy.\n"
    "Write a COMPLETE, ORIGINAL SKY Academy voiceover script on the given topic.\n"
    "CRITICAL RULE 1: telugu_text must contain ZERO emoji characters.\n"
    "CRITICAL RULE 2: ALL numbers in telugu_text must be written as English words (nineteen forty-seven, not 1947).\n"
    "CRITICAL RULE 3: ALL Telugu words must be in Telugu Unicode script. NEVER Roman transliteration.\n"
    "  WRONG: 'eeroju manam chuddam ookenaa' | CORRECT: 'ఈరోజు మనం చూద్దాం ఓకేనా'\n"
    "CRITICAL RULE 4: Each segment MUST be 150-180 words. Do not deviate from this range.\n"
    "CRITICAL RULE 5: Output must be a complete, valid JSON array -- do not truncate or stop mid-output.\n"
) + _SKY_DNA + _OUTPUT_FORMAT

_SYSTEM_TRANSCRIPT = (
    "You are an expert Telugu video script writer for SKY Academy.\n"
    "TRANSFORM competitor transcripts into a single 100% original SKY Academy script.\n"
    "CRITICAL RULE 1: telugu_text must contain ZERO emoji characters.\n"
    "CRITICAL RULE 2: ALL numbers in telugu_text must be written as English words.\n"
    "CRITICAL RULE 3: ALL Telugu words must be in Telugu Unicode script. NEVER Roman transliteration.\n"
    "  WRONG: 'eeroju manam chuddam ookenaa' | CORRECT: 'ఈరోజు మనం చూద్దాం ఓకేనా'\n"
    "CRITICAL RULE 4: Each segment MUST be 150-180 words.\n"
    "CRITICAL RULE 5: Output must be a complete, valid JSON array.\n\n"
    "CASE A -- Same topic: SYNTHESIZE all data.\n"
    "CASE B -- Different aspects: MERGE and INTERLINK.\n"
    "CASE C -- Redundant: Take BEST from each, enrich.\n"
    "Step 1 -- STRIP ALL COMPETITOR TRACES\n"
    "Step 2 -- ENRICH CONTENT (+25% minimum)\n"
    "Step 3 -- FULL SKY ACADEMY VOICE IN TELUGU LIPI\n"
) + _SKY_DNA + _OUTPUT_FORMAT

_SYSTEM_PDF = (
    "You are an expert Telugu video script writer for SKY Academy.\n"
    "CONVERT dry book/study material into an engaging SKY Academy voiceover.\n"
    "CRITICAL RULE 1: telugu_text must contain ZERO emoji characters.\n"
    "CRITICAL RULE 2: ALL numbers in telugu_text must be written as English words.\n"
    "CRITICAL RULE 3: ALL Telugu words must be in Telugu Unicode script. NEVER Roman transliteration.\n"
    "  WRONG: 'idi chala important' | CORRECT: 'ఇది చాలా important'\n"
    "CRITICAL RULE 4: Each segment MUST be 150-180 words.\n"
    "CRITICAL RULE 5: Output must be a complete, valid JSON array.\n"
    "Step 1 -- DESTROY THE BOOKISH TONE\n"
    "Step 2 -- INJECT LIFE AND DEPTH\n"
    "Step 3 -- FULL SKY ACADEMY VOICE IN TELUGU LIPI\n"
) + _SKY_DNA + _OUTPUT_FORMAT


# ============================================================
# UTILITY -- STRIP EMOJIS
# ============================================================
def strip_emojis(text: str) -> str:
    emoji_pattern = re.compile(
        "["
        "\U0001F1E0-\U0001F1FF\U0001F300-\U0001F5FF\U0001F600-\U0001F64F"
        "\U0001F680-\U0001F6FF\U0001F700-\U0001F77F\U0001F780-\U0001F7FF"
        "\U0001F800-\U0001F8FF\U0001F900-\U0001F9FF\U0001FA00-\U0001FA6F"
        "\U0001FA70-\U0001FAFF\u2600-\u26FF\u2700-\u27BF\uFE0F\u200D"
        "\u23CF\u23E9-\u23F3\u231A-\u231B\u25AA-\u25FE\u2614-\u2615"
        "\u2648-\u2653\u267F\u2693\u26A1\u26AA-\u26AB\u26BD-\u26BE"
        "\u26C4-\u26C5\u26CE\u26D4\u26EA\u26F2-\u26F3\u26F5\u26FA\u26FD"
        "]+",
        flags=re.UNICODE,
    )
    return emoji_pattern.sub("", text).strip()


def store_new_chunks(parsed: list):
    for i, c in enumerate(parsed):
        cleaned = strip_emojis(c.get("telugu_text", ""))
        c["telugu_text"] = cleaned
        st.session_state[f"tv_{i}"] = cleaned
    st.session_state.chunks   = parsed
    st.session_state.seo_pack = None


def sync_edits_to_chunks() -> list:
    chunks = st.session_state.chunks
    if not chunks:
        return []
    synced = []
    for i, c in enumerate(chunks):
        synced.append({
            "seg":         c.get("seg", i + 1),
            "title":       c.get("title", f"Segment {i+1}"),
            "telugu_text": strip_emojis(
                st.session_state.get(f"tv_{i}", c.get("telugu_text", ""))
            ),
        })
    return synced


# ============================================================
# ROBUST JSON PARSING UTILITIES
# ============================================================

def _sanitize_json_str(s: str) -> str:
    result = []
    in_string = False
    escape_next = False
    for c in s:
        if escape_next:
            result.append(c)
            escape_next = False
            continue
        if c == '\\' and in_string:
            result.append(c)
            escape_next = True
            continue
        if c == '"':
            in_string = not in_string
            result.append(c)
            continue
        if in_string:
            if c == '\n':
                result.append('\\n')
            elif c == '\r':
                result.append('\\r')
            elif c == '\t':
                result.append('\\t')
            elif ord(c) < 32:
                result.append(f'\\u{ord(c):04x}')
            else:
                result.append(c)
        else:
            result.append(c)
    return ''.join(result)


def _extract_complete_objects(json_str: str) -> list:
    segments = []
    i = 0
    n = len(json_str)

    while i < n:
        start = json_str.find('{', i)
        if start == -1:
            break

        depth = 0
        in_string = False
        escape_next = False
        j = start
        found_end = -1

        while j < n:
            c = json_str[j]
            if escape_next:
                escape_next = False
            elif c == '\\' and in_string:
                escape_next = True
            elif c == '"':
                in_string = not in_string
            elif not in_string:
                if c == '{':
                    depth += 1
                elif c == '}':
                    depth -= 1
                    if depth == 0:
                        found_end = j
                        break
            j += 1

        if found_end == -1:
            break

        obj_str = json_str[start:found_end + 1]
        parsed_obj = None
        for attempt_str in [obj_str, _sanitize_json_str(obj_str)]:
            try:
                parsed_obj = json.loads(attempt_str)
                if isinstance(parsed_obj, dict):
                    break
            except Exception:
                parsed_obj = None

        if parsed_obj and isinstance(parsed_obj, dict):
            if 'telugu_text' in parsed_obj or 'seg' in parsed_obj:
                if 'seg' not in parsed_obj:
                    parsed_obj['seg'] = len(segments) + 1
                if 'telugu_text' not in parsed_obj:
                    parsed_obj['telugu_text'] = ''
                if 'title' not in parsed_obj:
                    parsed_obj['title'] = f"Segment {parsed_obj.get('seg', len(segments)+1)}"
                segments.append(parsed_obj)

        i = found_end + 1

    return segments


def parse_segments(raw: str):
    if not raw or not raw.strip():
        return None

    cleaned = raw.strip()
    cleaned = re.sub(r"^```[a-zA-Z]*\s*\n?", "", cleaned)
    cleaned = re.sub(r"\n?```\s*$", "", cleaned).strip()

    for attempt in [cleaned, _sanitize_json_str(cleaned)]:
        try:
            result = json.loads(attempt)
            if isinstance(result, list) and result:
                return result
            if isinstance(result, dict) and ('telugu_text' in result or 'seg' in result):
                return [result]
        except Exception:
            pass

    start_idx = cleaned.find("[")
    end_idx   = cleaned.rfind("]")
    if start_idx != -1 and end_idx > start_idx:
        chunk = cleaned[start_idx:end_idx + 1]
        for attempt in [chunk, _sanitize_json_str(chunk)]:
            try:
                result = json.loads(attempt)
                if isinstance(result, list) and result:
                    return result
            except Exception:
                pass

    search_start = start_idx if start_idx != -1 else 0
    segments = _extract_complete_objects(cleaned[search_start:])
    if segments:
        return sorted(segments, key=lambda x: x.get('seg', 999))

    if search_start > 0:
        segments = _extract_complete_objects(cleaned)
        if segments:
            return sorted(segments, key=lambda x: x.get('seg', 999))

    for attempt in [raw, re.sub(r"```(?:json)?", "", raw).strip()]:
        try:
            result = json.loads(attempt)
            if isinstance(result, list) and result:
                return result
        except Exception:
            pass
        m = re.search(r"\[[\s\S]*\]", attempt)
        if m:
            for s in [m.group(), _sanitize_json_str(m.group())]:
                try:
                    result = json.loads(s)
                    if isinstance(result, list) and result:
                        return result
                except Exception:
                    pass

    return None


# ============================================================
# PROMPT BUILDERS
# ============================================================
def _inject_counts(system: str, num_segs: int) -> str:
    return (
        system
        .replace("{NUM_SEGS}", str(num_segs))
        .replace("{WORDS_PER_SEGMENT}", str(WORDS_PER_SEGMENT))
    )


def build_prompts_topic(topic, num_segs, special_instructions, video_type="subjective"):
    vdna   = _DNA_GENERAL_VIDEO if video_type == "general" else _DNA_SUBJECTIVE_VIDEO
    system = _inject_counts(_SYSTEM_TOPIC + "\n\n" + TELUGU_TTS_MASTER_PROMPT + "\n\n" + vdna, num_segs)
    si     = f"\nSpecial Instructions:\n{special_instructions.strip()}" if special_instructions.strip() else ""
    user   = (
        f"Generate a complete SKY Academy Telugu video script on:\n\n"
        f"**Topic:** {topic.strip()}\n"
        f"**Video Type:** {'General/Strategy/Motivation' if video_type=='general' else 'Subjective/Deep Teaching'}\n"
        f"**Segments required:** {num_segs}\n"
        f"**Words per segment:** STRICTLY 150-180 words -- count carefully and stay within range\n{si}\n\n"
        f"CRITICAL REMINDERS:\n"
        f"- ALL Telugu words MUST be in Telugu Unicode script (ఇలా), NEVER Roman (like this)\n"
        f"- Segment 1: HOOK (Telugu lipi) -> Welcome -> Content\n"
        f"- Last segment: SKY Academy CTA in Telugu script\n"
        f"- ZERO emojis in telugu_text\n"
        f"- ALL numbers as English words (nineteen forty-seven, forty-second Amendment)\n"
        f"- Lists/series: humanized flow with connectors, NOT robotic listing\n"
        f"- Output must be a COMPLETE valid JSON array ending with ]\n"
        f"- Return ONLY valid JSON array, no other text"
    )
    return system, user


def build_prompts_multi_transcript(transcripts, topic_hint, num_segs,
                                   special_instructions, merge_mode="auto", video_type="subjective"):
    vdna   = _DNA_GENERAL_VIDEO if video_type == "general" else _DNA_SUBJECTIVE_VIDEO
    system = _inject_counts(_SYSTEM_TRANSCRIPT + "\n\n" + TELUGU_TTS_MASTER_PROMPT + "\n\n" + vdna, num_segs)
    n      = len(transcripts)
    merge_guide = {
        "auto":          "AUTO-DETECT the relationship and apply CASE A/B/C.",
        "synthesize":    "SYNTHESIZE -- same topic, different data.",
        "merge_aspects": "MERGE ASPECTS -- different topics, one narrative.",
    }[merge_mode]
    hint = f"\n**Topic context:** {topic_hint.strip()}" if topic_hint.strip() else ""
    si   = f"\nSpecial Instructions:\n{special_instructions.strip()}" if special_instructions.strip() else ""
    transcript_blocks = "\n\n".join(
        f"COMPETITOR TRANSCRIPT {i+1} of {n}\n[File: {t['filename']}]\n{t['text'].strip()}\nEND TRANSCRIPT {i+1}"
        for i, t in enumerate(transcripts)
    )
    user = (
        f"Transform {n} transcript(s) into a SINGLE SKY Academy voiceover.\n"
        f"**Video Type:** {'General/Strategy' if video_type=='general' else 'Subjective/Teaching'}\n"
        f"**Segments required:** {num_segs}\n"
        f"**Words per segment:** STRICTLY 150-180 words\n"
        f"**Merge strategy:** {merge_guide}\n{hint}{si}\n\n"
        f"{transcript_blocks}\n\n"
        f"CRITICAL REMINDERS:\n"
        f"- ALL Telugu words in Telugu Unicode script (ఇలా), NEVER Roman\n"
        f"- Strip ALL competitor traces\n- Add 25%+ more value\n"
        f"- Last segment: SKY Academy CTA\n- ZERO emojis in telugu_text\n"
        f"- ALL numbers as English words\n- Lists: humanized flow, not robotic\n"
        f"- Output must be a COMPLETE valid JSON array ending with ]\n"
        f"- Return ONLY valid JSON"
    )
    return system, user


def build_prompts_pdf(pdf_text, topic_hint, num_segs, special_instructions, video_type="subjective"):
    vdna   = _DNA_GENERAL_VIDEO if video_type == "general" else _DNA_SUBJECTIVE_VIDEO
    system = _inject_counts(_SYSTEM_PDF + "\n\n" + TELUGU_TTS_MASTER_PROMPT + "\n\n" + vdna, num_segs)
    hint   = f"\n**Topic/Chapter context:** {topic_hint.strip()}" if topic_hint.strip() else ""
    si     = f"\nSpecial Instructions:\n{special_instructions.strip()}" if special_instructions.strip() else ""
    user   = (
        f"Convert book/study material into SKY Academy voiceover.\n"
        f"**Video Type:** {'General/Strategy' if video_type=='general' else 'Subjective/Teaching'}\n"
        f"**Segments required:** {num_segs}\n"
        f"**Words per segment:** STRICTLY 150-180 words\n{hint}{si}\n\n"
        f"BOOK/STUDY MATERIAL:\n{pdf_text.strip()}\n\n"
        f"CRITICAL REMINDERS:\n"
        f"- ALL Telugu words in Telugu Unicode script (ఇలా), NEVER Roman\n"
        f"- Destroy bookish tone\n- Last segment: SKY Academy CTA\n"
        f"- ZERO emojis in telugu_text\n- ALL numbers as English words\n"
        f"- Lists: humanized flow with connectors\n"
        f"- Output must be a COMPLETE valid JSON array ending with ]\n"
        f"- Return ONLY valid JSON"
    )
    return system, user


def build_regen_segment_prompt(video_type, topic, chunks, seg_idx, instruction, num_segs):
    vdna     = _DNA_GENERAL_VIDEO if video_type == "general" else _DNA_SUBJECTIVE_VIDEO
    is_first = seg_idx == 0
    is_last  = seg_idx == len(chunks) - 1
    prev_text = chunks[seg_idx-1].get("telugu_text","")[:350] if seg_idx > 0 else ""
    next_text = chunks[seg_idx+1].get("telugu_text","")[:350] if seg_idx < len(chunks)-1 else ""
    system = (
        "You are an expert Telugu video script writer for SKY Academy.\n"
        "Regenerate EXACTLY ONE segment based on a specific instruction.\n"
        "CRITICAL: telugu_text MUST contain ZERO emoji characters.\n"
        "CRITICAL: ALL numbers must be written as English words.\n"
        "CRITICAL: ALL Telugu words must be in Telugu Unicode script. NEVER Roman transliteration.\n"
        "CRITICAL: Segment must be 150-180 words.\n\n"
    ) + vdna + _SKY_DNA
    user = (
        f"Topic: {topic}\nTotal segments: {num_segs}\nRegenerating: Segment {seg_idx+1}\n"
        f"Instruction: {instruction or 'Improve -- better flow, sharper memory hints'}\n\n"
        + (f"PREVIOUS SEGMENT:\n{prev_text}\n\n" if prev_text else "")
        + f"CURRENT SEGMENT:\n{chunks[seg_idx].get('telugu_text','')}\n\n"
        + (f"NEXT SEGMENT:\n{next_text}\n\n" if next_text else "")
        + "RULES:\n"
        + ("- Segment 1: Hook + Welcome + Content\n" if is_first else "")
        + ("- LAST segment: SKY Academy CTA + Telegram CTA\n" if is_last else "")
        + f"- STRICTLY 150-180 words\n- ZERO emojis\n"
        + f"- ALL Telugu words in Telugu Unicode script\n"
        + f"- ALL numbers as English words\n"
        + f"- Lists/series: humanized flow with connectors\n"
        + f"- Return ONLY valid JSON for ONE segment:\n"
        + '{"seg":' + str(seg_idx+1) + ',"title":"3-5 word English heading","telugu_text":"...Telugu Unicode script..."}'
    )
    return system, user


def build_seo_prompt(topic, chunks, video_type):
    preview = " ".join(c.get("telugu_text","")[:180] for c in chunks[:3])
    system  = (
        "You are a YouTube SEO expert for Telugu competitive exam content. "
        "Return ONLY valid JSON, no markdown."
    )
    user = (
        f"Script Topic: {topic}\n"
        f"Video Type: {'General/Strategy' if video_type=='general' else 'Deep Teaching'}\n"
        f"Script Preview: {preview[:450]}\n\n"
        f"Generate:\n1. ONE YouTube TITLE (max 70 chars, Telugu+English mix)\n"
        f"2. 20 TAGS (comma-separated)\n"
        f"Return ONLY: {{\"title\": \"...\", \"tags\": \"tag1, tag2, ...\"}}"
    )
    return system, user


def build_thumbnail_suggest_prompt(topic, chunks, video_type):
    preview = ""
    if chunks:
        preview = " ".join(c.get("telugu_text", "")[:120] for c in chunks[:3])
    system = (
        "You are a YouTube thumbnail CTR optimization expert for Indian competitive exam channels. "
        "Return ONLY valid JSON, no markdown, no extra text."
    )
    user = f"""
Generate high-CTR YouTube thumbnail text for SKY Academy Telugu competitive exam channel.

Topic: {topic or "Competitive Exam Preparation"}
Video Type: {"General Strategy/Motivation" if video_type == "general" else "Deep Subject Teaching"}
Script Preview: {preview[:400]}

Generate EXACTLY 4 lines optimized for maximum click-through rate:

Line 1 (line1): Target audience / exam name — short, bold, uppercase
  Examples: "AP HIGH COURT 2026", "SSC CGL 2026", "APPSC GROUP 1", "AP TET 2026", "UPSC 2026"

Line 2 (line2): Core video topic — what makes them MUST-click
  Examples: "POLITY HIGH YIELD", "BNS COMPLETE SERIES", "GEOGRAPHY STRATEGY", "CURRENT AFFAIRS BLAST"

Line 3 (line3): Brief compelling sub-explanation
  Examples: "WITH PYQ ANALYSIS", "EXAM GUARANTEED TOPICS", "STRATEGY + PDF FREE", "100 MARKS SHORTCUT"

Line 4 (line4): Punchy Telugu text (MUST be in Telugu Unicode script)
  Examples: "Miss అవ్వాండి!", "ఇప్పుడే చూడండి!", "ఒక్కసారి చూస్తే చాలు!", "ఇది మీకోసమే!"

Return ONLY valid JSON:
{{"line1": "...", "line2": "...", "line3": "...", "line4": "..."}}
"""
    return system, user


# ============================================================
# AI HANDOUT PROMPT BUILDERS
# ============================================================
def build_handout_prompt(topic: str, chunks: list, num_pages: int) -> tuple:
    script_outline = ""
    for i, c in enumerate(chunks[:10]):
        title = c.get("title", f"Segment {i+1}")
        script_outline += f"Segment {i+1}: {title}\n"

    target_words = num_pages * 380

    system = (
        "You are an expert educational content creator for SKY Academy. "
        "Create comprehensive, accurate, print-ready study handouts in ENGLISH ONLY. "
        "Include real data, dates, statistics, article numbers, case names. "
        "Return ONLY valid JSON. No markdown fences. No text outside the JSON."
    )

    user = f"""Create a comprehensive study handout for competitive exam students.

Topic: {topic}
Target: approximately {num_pages} A4 printed page(s) (~{target_words} words of content)

Script outline:
{script_outline}

Generate detailed JSON with this EXACT structure:
{{
  "topic_title": "Full descriptive topic title",
  "exam_importance": "2-3 sentences about why this topic appears in UPSC/APPSC/TSPSC/SSC exams",
  "key_facts": [
    {{"label": "short label", "value": "specific fact/date/data"}},
    {{"label": "short label", "value": "specific fact/date/data"}}
  ],
  "sections": [
    {{
      "heading": "Introduction / Historical Background",
      "text": "Write 2-3 full paragraphs here with actual context, history, and significance.",
      "bullets": [],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Key Concepts / Important Terms",
      "text": "1-2 paragraph explanation with real data.",
      "bullets": ["specific point with actual data", "another specific point"],
      "table": {{
        "headers": ["Term / Item", "Meaning / Description", "Significance / Example"],
        "rows": [
          ["term1", "meaning1", "significance1"],
          ["term2", "meaning2", "significance2"],
          ["term3", "meaning3", "significance3"],
          ["term4", "meaning4", "significance4"]
        ]
      }},
      "questions": []
    }},
    {{
      "heading": "Important Data / Current Affairs",
      "text": "Explanation paragraph with latest data.",
      "bullets": ["data point 1 with specific numbers", "data point 2"],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Memory Tricks & Quick Connections",
      "text": "Use these clever connections to never forget these facts:",
      "bullets": ["Memory trick 1 -- explain the connection clearly", "Memory trick 2 -- fun and memorable"],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Previous Year Questions (PYQ Style)",
      "text": "",
      "bullets": [],
      "table": null,
      "questions": [
        "Q1. Full question text here? (UPSC/APPSC Year)",
        "Q2. Full question text here? (TSPSC Year)",
        "Q3. Full question text here? (SSC Year)"
      ]
    }}
  ]
}}

STRICT RULES:
1. ENGLISH ONLY -- absolutely no Telugu text
2. Every "text" field MUST have 2-3 full sentences minimum
3. Include REAL specific data: exact dates, article numbers, amendment numbers
4. Tables must have minimum 4 rows of meaningful data
5. Memory tricks must be instantly obvious and clever
6. PYQs must be realistic exam-quality questions
7. Generate enough content to fill approximately {num_pages} A4 pages
8. Return ONLY valid JSON -- nothing else"""

    return system, user


def build_standalone_handout_prompt(topic_overview: str, num_pages: int) -> tuple:
    target_words = num_pages * 380

    system = (
        "You are an expert educational content creator for SKY Academy. "
        "Create comprehensive, accurate, print-ready study handouts in ENGLISH ONLY. "
        "Include real data, dates, statistics, article numbers, case names. "
        "Return ONLY valid JSON. No markdown fences. No text outside the JSON."
    )

    user = f"""Create a comprehensive study handout for competitive exam students based on the following topic overview and directions.

TOPIC OVERVIEW & DIRECTIONS:
{topic_overview.strip()}

Target: approximately {num_pages} A4 printed page(s) (~{target_words} words of content)

Generate detailed JSON with this EXACT structure:
{{
  "topic_title": "Full descriptive topic title",
  "exam_importance": "2-3 sentences about why this topic appears in UPSC/APPSC/TSPSC/SSC exams",
  "key_facts": [
    {{"label": "short label", "value": "specific fact/date/data"}},
    {{"label": "short label", "value": "specific fact/date/data"}}
  ],
  "sections": [
    {{
      "heading": "Introduction / Background",
      "text": "Write 2-3 full paragraphs with actual context, history, and significance.",
      "bullets": [],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Core Concepts & Key Terms",
      "text": "1-2 paragraph explanation with real data.",
      "bullets": ["specific point", "another specific point"],
      "table": {{
        "headers": ["Item", "Description", "Significance"],
        "rows": [
          ["item1", "desc1", "sig1"],
          ["item2", "desc2", "sig2"],
          ["item3", "desc3", "sig3"],
          ["item4", "desc4", "sig4"]
        ]
      }},
      "questions": []
    }},
    {{
      "heading": "Important Data & Statistics",
      "text": "Paragraph with key numbers and current data.",
      "bullets": ["data point 1", "data point 2"],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Memory Tricks & Quick Connections",
      "text": "Clever connections to never forget key facts:",
      "bullets": ["trick 1", "trick 2"],
      "table": null,
      "questions": []
    }},
    {{
      "heading": "Previous Year Questions (PYQ Style)",
      "text": "",
      "bullets": [],
      "table": null,
      "questions": [
        "Q1. Full question? (Exam Year)",
        "Q2. Full question? (Exam Year)",
        "Q3. Full question? (Exam Year)"
      ]
    }}
  ]
}}

STRICT RULES:
1. ENGLISH ONLY
2. Every "text" field: minimum 2-3 full sentences
3. Real specific data: exact dates, article numbers, statistics
4. Tables: minimum 4 rows
5. PYQs: realistic exam-quality questions with exam name and year
6. Fill approximately {num_pages} A4 pages
7. Return ONLY valid JSON"""

    return system, user


# ============================================================
# STUDY NOTES HTML (updated -- no slide_prompt)
# ============================================================
def generate_study_notes_html(chunks: list, video_type: str,
                               heading: str = "", youtube_link: str = "") -> str:
    accent_colors = ["#3b82f6","#ef4444","#16a34a","#f97316",
                     "#7c3aed","#0891b2","#be185d","#b45309"]

    sections_html = ""
    for i, chunk in enumerate(chunks):
        title      = chunk.get("title", f"Segment {i+1}")
        color      = accent_colors[i % len(accent_colors)]
        tel_text   = chunk.get("telugu_text","")
        preview    = tel_text[:280] + ("..." if len(tel_text) > 280 else "")

        sections_html += f"""
        <div class="ns" style="border-left-color:{color}">
          <div class="ns-head">
            <span class="ns-num" style="background:{color}">{i+1}</span>
            <span class="ns-title">{title}</span>
          </div>
          <p class="nb-text">{preview}</p>
          <div class="ann">
            <div class="ann-label">My Notes</div>
            <div class="ann-line"></div>
            <div class="ann-line"></div>
          </div>
        </div>"""

    if youtube_link.strip():
        vid_html = (
            f'<div class="intro">To better understand this PDF, click here: '
            f'<a href="{youtube_link.strip()}">{youtube_link.strip()}</a>'
            f' &nbsp;·&nbsp; <strong>SKY Academy Competitive Exam Preparation</strong></div>'
        )
    else:
        vid_html = (
            '<div class="intro">Read these notes while watching the video for maximum retention'
            ' &nbsp;·&nbsp; <strong>SKY Academy Competitive Exam Preparation</strong></div>'
        )

    heading_html = ""
    if heading.strip():
        heading_html = f'<div class="custom-heading">{heading.strip()}</div>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>SKY Academy -- Classroom Notes</title>
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',Arial,sans-serif;background:#dde3f0;padding:20px;color:#111;font-size:14px}}
.page{{max-width:800px;margin:0 auto;background:#fff;border-radius:14px;overflow:hidden;box-shadow:0 6px 32px rgba(0,0,0,.15)}}
.hdr{{background:linear-gradient(135deg,#020024 0%,#090979 50%,#00d4ff 100%);padding:28px 26px 22px;text-align:center}}
.logo{{font-size:44px;font-weight:900;letter-spacing:10px;line-height:1;margin-bottom:3px}}
.ls{{color:#FF6B6B}}.lk{{color:#FFE66D}}.ly{{color:#4ECDC4}}
.acad{{color:rgba(255,255,255,.7);font-size:9px;font-weight:800;letter-spacing:5px;text-transform:uppercase;margin-bottom:10px}}
.custom-heading{{color:#FFE66D;font-size:20px;font-weight:800;margin-top:8px;line-height:1.3}}
.nbadge{{margin-top:10px;background:rgba(255,255,255,.16);color:#fff;border-radius:20px;padding:4px 18px;font-size:10px;font-weight:700;letter-spacing:2px;display:inline-block;text-transform:uppercase}}
.intro{{background:#eff6ff;padding:9px 20px;text-align:center;font-size:12px;color:#1e40af;border-bottom:1px solid #bfdbfe}}
.intro a{{color:#1e40af;word-break:break-all}}
.pbtn-wrap{{text-align:center;padding:14px 20px 6px}}
.pbtn{{background:linear-gradient(135deg,#020024,#090979);color:#fff;border:none;border-radius:8px;padding:10px 30px;font-size:13px;font-weight:700;cursor:pointer;letter-spacing:1px}}
.phint{{font-size:11px;color:#9ca3af;margin-top:5px}}
.body{{padding:14px 20px 20px}}
.ns{{background:#fafafa;border:1px solid #e5e7eb;border-left:4px solid;border-radius:8px;margin-bottom:14px;overflow:hidden;page-break-inside:avoid}}
.ns-head{{display:flex;align-items:center;gap:9px;padding:9px 12px;background:rgba(0,0,0,.03);border-bottom:1px solid #e5e7eb}}
.ns-num{{width:24px;height:24px;min-width:24px;border-radius:50%;color:#fff;font-size:11px;font-weight:800;display:flex;align-items:center;justify-content:center}}
.ns-title{{font-size:13.5px;font-weight:700;color:#111827;line-height:1.3}}
.nb-text{{padding:8px 12px;font-size:12.5px;color:#374151;line-height:1.75;font-family:'Noto Sans Telugu',sans-serif}}
.ann{{padding:7px 12px 9px;border-top:1px dashed #e5e7eb}}
.ann-label{{font-size:9px;color:#bbb;font-weight:600;letter-spacing:1px;text-transform:uppercase;margin-bottom:4px}}
.ann-line{{height:1px;background:#ebebeb;margin-bottom:8px}}
.tg{{margin:4px 20px 20px;background:linear-gradient(to right,#eff6ff,#dbeafe);border:2px solid #3b82f6;border-radius:12px;padding:16px 20px;text-align:center}}
.tg h4{{color:#1e40af;font-size:14px;font-weight:800;margin-bottom:7px}}
.tg p{{color:#1e40af;font-size:12.5px;line-height:1.85}}
.ftr{{background:linear-gradient(135deg,#020024,#090979);padding:13px 20px;text-align:center}}
.flogo{{font-size:18px;font-weight:900;letter-spacing:5px;margin-bottom:3px}}
.fsub{{color:rgba(255,255,255,.45);font-size:9px}}
@media print{{
  body{{background:#fff;padding:0}}
  .page{{box-shadow:none;border-radius:0;max-width:100%}}
  .pbtn-wrap,.tg{{display:none!important}}
  .ns{{break-inside:avoid}}
  @page{{margin:1.2cm 1.5cm}}
}}
</style>
</head>
<body>
<div class="page">
  <div class="hdr">
    <div class="logo"><span class="ls">S</span><span class="lk">K</span><span class="ly">Y</span></div>
    <div class="acad">Academy</div>
    {heading_html}
    <div><span class="nbadge">Classroom Notes</span></div>
  </div>
  {vid_html}
  <div class="pbtn-wrap">
    <button class="pbtn" onclick="window.print()">Print / Save as PDF</button>
    <p class="phint">Use Chrome or Edge for best results</p>
  </div>
  <div class="body">{sections_html}</div>
  <div class="tg">
    <h4>SKY Academy Telegram Channel</h4>
    <p>Free notes, daily PDFs, previous year questions and exam alerts --<br>
    Join and take your preparation to the next level!</p>
  </div>
  <div class="ftr">
    <div class="flogo">
      <span style="color:#FF6B6B">S</span>
      <span style="color:#FFE66D">K</span>
      <span style="color:#4ECDC4">Y</span>
    </div>
    <div class="fsub">SKY ACADEMY &nbsp;|&nbsp; Script Engine v4.0 &nbsp;|&nbsp; Internal Tool</div>
  </div>
</div>
</body>
</html>"""


# ============================================================
# PARSE HANDOUT JSON
# ============================================================
def parse_handout_content(raw: str):
    cleaned = raw.strip()
    cleaned = re.sub(r"^```[a-zA-Z]*\s*\n?", "", cleaned)
    cleaned = re.sub(r"\n?```\s*$", "", cleaned).strip()
    for attempt in [cleaned, _sanitize_json_str(cleaned)]:
        try:
            result = json.loads(attempt)
            if isinstance(result, dict):
                return result
        except Exception:
            pass
    m = re.search(r'\{[\s\S]*\}', cleaned)
    if m:
        for attempt in [m.group(), _sanitize_json_str(m.group())]:
            try:
                result = json.loads(attempt)
                if isinstance(result, dict):
                    return result
            except Exception:
                pass
    return None


# ============================================================
# RENDER AI HANDOUT HTML
# ============================================================
def render_handout_html(heading: str, youtube_link: str,
                        content_data: dict, topic: str) -> str:
    topic_title     = content_data.get("topic_title", heading or topic)
    exam_importance = content_data.get("exam_importance", "")
    key_facts       = content_data.get("key_facts", [])
    sections        = content_data.get("sections", [])

    display_heading = heading.strip() if heading.strip() else topic_title

    if youtube_link.strip():
        vid_html = (
            f'<div class="vid-banner">To better understand this PDF, click here: '
            f'<a href="{youtube_link.strip()}">{youtube_link.strip()}</a></div>'
        )
    else:
        vid_html = (
            '<div class="vid-banner">For better understanding, please watch our '
            '<a href="https://www.youtube.com/@Skyacademytelugu">'
            'https://www.youtube.com/@Skyacademytelugu</a> YouTube channel.</div>'
        )

    importance_html = ""
    if exam_importance:
        importance_html = (
            f'<div class="exam-imp">'
            f'<span class="eit">Why it Matters for Exams: </span>'
            f'<span class="eit-text">{exam_importance}</span></div>'
        )

    facts_html = ""
    if key_facts:
        facts_html = '<div class="kf-grid">'
        for kf in key_facts:
            facts_html += (
                f'<div class="kf-box">'
                f'<div class="kf-lbl">{kf.get("label","")}</div>'
                f'<div class="kf-val">{kf.get("value","")}</div>'
                f'</div>'
            )
        facts_html += '</div>'

    sec_colors = ["#1e293b","#1e3a5f","#1a3a2a","#2d1b4e","#3b1a1a",
                  "#1a2d3b","#2d2000","#0d2d1a"]
    sections_html = ""
    for idx, sec in enumerate(sections):
        sh        = sec.get("heading","")
        text      = sec.get("text","")
        bullets   = sec.get("bullets",[])
        tbl       = sec.get("table",None)
        questions = sec.get("questions",[])
        hdr_color = sec_colors[idx % len(sec_colors)]

        table_html = ""
        if tbl and isinstance(tbl, dict):
            headers = tbl.get("headers",[])
            rows    = tbl.get("rows",[])
            if headers and rows:
                table_html = "<table class='htbl'><thead><tr>"
                for h in headers:
                    table_html += f"<th>{h}</th>"
                table_html += "</tr></thead><tbody>"
                for ri, row in enumerate(rows):
                    rc = "even" if ri % 2 == 0 else ""
                    table_html += f"<tr class='{rc}'>"
                    for cell in row:
                        table_html += f"<td>{cell}</td>"
                    table_html += "</tr>"
                table_html += "</tbody></table>"

        bullets_html = ""
        if bullets:
            bullets_html = (
                "<ul class='hbul'>"
                + "".join(f"<li>{b}</li>" for b in bullets)
                + "</ul>"
            )

        questions_html = ""
        if questions:
            questions_html = (
                "<ol class='hpq'>"
                + "".join(f"<li>{q}</li>" for q in questions)
                + "</ol>"
            )

        sections_html += f"""
        <div class="sbox">
          <div class="shdr" style="background:{hdr_color}">&#9632; {sh}</div>
          <div class="sbody">
            {"<p>" + text + "</p>" if text else ""}
            {table_html}
            {bullets_html}
            {questions_html}
          </div>
        </div>"""

    footer_html = (
        '<div class="pg-footer">'
        'For more videos and free study materials, Join our app: '
        '<a href="https://shorturl.at/rpTqP">https://shorturl.at/rpTqP</a>'
        '&nbsp;&nbsp;|&nbsp;&nbsp;'
        'For one-to-one Mentorship: Send <strong>Hi</strong> on WhatsApp '
        '<strong>9491370061</strong>'
        '</div>'
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>SKY Academy -- {display_heading}</title>
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',Arial,sans-serif;background:#d4d8e8;padding:20px;color:#111;font-size:13px;line-height:1.65;padding-bottom:60px}}
.page{{max-width:820px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 6px 32px rgba(0,0,0,.18);margin-bottom:60px}}
.pbtn-wrap{{text-align:center;padding:14px 20px 6px;background:#f8f9ff}}
.pbtn{{background:linear-gradient(135deg,#020024,#090979);color:#fff;border:none;border-radius:8px;padding:10px 32px;font-size:13px;font-weight:700;cursor:pointer;letter-spacing:1px}}
.phint{{font-size:11px;color:#9ca3af;margin-top:5px}}
.hdr{{background:linear-gradient(135deg,#020024 0%,#090979 50%,#00d4ff 100%);padding:24px 26px 20px;text-align:center}}
.logo{{font-size:42px;font-weight:900;letter-spacing:10px;line-height:1;margin-bottom:2px}}
.ls{{color:#FF6B6B}}.lk{{color:#FFE66D}}.ly{{color:#4ECDC4}}
.acad{{color:rgba(255,255,255,.65);font-size:8px;font-weight:800;letter-spacing:5px;text-transform:uppercase;margin-bottom:10px}}
.htitle{{color:#FFE66D;font-size:22px;font-weight:800;line-height:1.3;margin:0 auto;max-width:92%}}
.hsubtitle{{color:rgba(255,255,255,.5);font-size:9px;font-weight:600;letter-spacing:2px;text-transform:uppercase;margin-top:6px}}
.vid-banner{{background:#1e3a5f;color:#93c5fd;padding:8px 20px;font-size:11.5px;text-align:center;border-bottom:1px solid #1d4ed8}}
.vid-banner a{{color:#60a5fa;word-break:break-all}}
.exam-imp{{background:#fffbeb;border:1px solid #fde68a;border-left:4px solid #f59e0b;margin:14px 16px 10px;padding:10px 14px;border-radius:4px;font-size:12.5px;font-style:italic;color:#78350f}}
.eit{{font-weight:800;font-style:normal}}
.kf-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(175px,1fr));gap:8px;margin:10px 16px 14px}}
.kf-box{{background:#1e293b;border-radius:6px;padding:8px 12px}}
.kf-lbl{{font-size:9px;font-weight:700;letter-spacing:1px;text-transform:uppercase;color:#94a3b8;margin-bottom:3px}}
.kf-val{{font-size:12px;font-weight:700;color:#e2e8f0;line-height:1.4}}
.sbox{{margin:0 16px 12px;border-radius:6px;overflow:hidden;border:1px solid #e5e7eb;page-break-inside:avoid}}
.shdr{{padding:9px 14px;font-size:13px;font-weight:800;color:#fff;letter-spacing:0.3px}}
.sbody{{padding:12px 14px;background:#fff}}
.sbody p{{font-size:13px;color:#374151;line-height:1.8;margin-bottom:8px}}
.htbl{{width:100%;border-collapse:collapse;margin:10px 0;font-size:12px}}
.htbl th{{background:#1e293b;color:#fff;padding:8px 10px;text-align:left;font-size:11px;font-weight:700}}
.htbl td{{padding:7px 10px;border-bottom:1px solid #e5e7eb;color:#374151;vertical-align:top;line-height:1.6}}
.htbl tr.even td{{background:#f8fafc}}
.htbl tr:last-child td{{border-bottom:none}}
.hbul{{padding-left:20px;margin:8px 0}}
.hbul li{{font-size:13px;color:#374151;line-height:1.85;margin-bottom:4px}}
.hbul li::marker{{color:#3b82f6;font-size:14px}}
.hpq{{padding-left:22px;margin:8px 0}}
.hpq li{{font-size:13px;color:#1e293b;line-height:1.85;margin-bottom:8px;font-weight:500}}
.pg-footer{{background:linear-gradient(135deg,#020024,#090979);color:rgba(255,255,255,.8);padding:10px 20px;font-size:10.5px;text-align:center;position:fixed;bottom:0;left:0;right:0;z-index:999}}
.pg-footer a{{color:#93c5fd}}
@media print{{
  body{{background:#fff;padding:0;padding-bottom:0}}
  .page{{box-shadow:none;border-radius:0;max-width:100%;margin-bottom:0}}
  .pbtn-wrap{{display:none!important}}
  .pg-footer{{position:fixed;bottom:0;left:0;right:0;font-size:9px;padding:6px 16px}}
  .sbox{{break-inside:avoid}}
  @page{{margin:1.2cm 1.5cm 2.2cm 1.5cm}}
}}
</style>
</head>
<body>
<div class="page">
  <div class="pbtn-wrap">
    <button class="pbtn" onclick="window.print()">Print / Save as PDF</button>
    <p class="phint">Chrome or Edge -- Ctrl+P -- Save as PDF</p>
  </div>
  <div class="hdr">
    <div class="logo"><span class="ls">S</span><span class="lk">K</span><span class="ly">Y</span></div>
    <div class="acad">Academy</div>
    <div class="htitle">{display_heading}</div>
    <div class="hsubtitle">Classroom Notes &nbsp;·&nbsp; Competitive Exam Preparation</div>
  </div>
  {vid_html}
  {importance_html}
  {facts_html}
  {sections_html}
</div>
{footer_html}
</body>
</html>"""


# ============================================================
# RENDER HANDOUT DOCX
# ============================================================
def render_handout_docx(heading: str, youtube_link: str,
                        content_data: dict, topic: str):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        t = doc.add_heading(heading.strip() if heading.strip() else
                            content_data.get("topic_title", topic), 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vl = doc.add_paragraph()
        if youtube_link.strip():
            vl.add_run(f"Video: {youtube_link.strip()}").italic = True
        else:
            vl.add_run("Watch: https://www.youtube.com/@Skyacademytelugu").italic = True
        vl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        ei = content_data.get("exam_importance","")
        if ei:
            doc.add_heading("Why it Matters for Exams", 2)
            doc.add_paragraph(ei)

        kf_list = content_data.get("key_facts",[])
        if kf_list:
            doc.add_heading("Key Facts", 2)
            for kf in kf_list:
                p = doc.add_paragraph()
                r = p.add_run(f"{kf.get('label','')}: ")
                r.bold = True
                p.add_run(kf.get("value",""))

        for sec in content_data.get("sections",[]):
            doc.add_heading(sec.get("heading",""), 2)
            text = sec.get("text","")
            if text:
                doc.add_paragraph(text)
            tbl = sec.get("table",None)
            if tbl and isinstance(tbl, dict):
                headers = tbl.get("headers",[])
                rows    = tbl.get("rows",[])
                if headers and rows:
                    table = doc.add_table(rows=1+len(rows), cols=len(headers))
                    table.style = "Table Grid"
                    hcells = table.rows[0].cells
                    for ci, h in enumerate(headers):
                        hcells[ci].text = h
                    for ri, row in enumerate(rows):
                        rcells = table.rows[ri+1].cells
                        for ci, cell in enumerate(row):
                            rcells[ci].text = str(cell)
            for b in sec.get("bullets",[]):
                doc.add_paragraph(b, style="List Bullet")
            for q in sec.get("questions",[]):
                doc.add_paragraph(q, style="List Number")

        doc.add_paragraph()
        fp = doc.add_paragraph(
            "For more videos and free study materials, Join our app: https://shorturl.at/rpTqP\n"
            "For one-to-one Mentorship: Send Hi on WhatsApp 9491370061"
        )
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None


# ============================================================
# FILE TEXT EXTRACTION
# ============================================================
def extract_pdf_text(file_bytes: bytes):
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages  = [p.extract_text() for p in reader.pages if p.extract_text()]
        if pages:
            return "\n".join(pages).strip(), f"pypdf · {len(reader.pages)} pages"
    except Exception:
        pass
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages  = [p.extract_text() for p in reader.pages if p.extract_text()]
        if pages:
            return "\n".join(pages).strip(), f"PyPDF2 · {len(reader.pages)} pages"
    except Exception:
        pass
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text()
                if t:
                    pages.append(t)
        if pages:
            return "\n".join(pages).strip(), f"pdfplumber · {len(pages)} pages"
    except Exception:
        pass
    return None, "No PDF library. Run: pip install pypdf"


def extract_docx_text(file_bytes: bytes):
    try:
        from docx import Document
        doc   = Document(io.BytesIO(file_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts)
        if text.strip():
            return text.strip(), f"python-docx · {len(doc.paragraphs)} paragraphs"
        return None, "DOCX file appears empty or image-only"
    except ImportError:
        return None, "python-docx not installed. Run: pip install python-docx"
    except Exception as exc:
        return None, f"DOCX error: {exc}"


def extract_any_file(file_bytes: bytes, filename: str):
    ext = filename.rsplit(".",1)[-1].lower() if "." in filename else ""
    if ext == "txt":
        for enc in ("utf-8","utf-16","latin-1"):
            try:
                text = file_bytes.decode(enc)
                if text.strip():
                    return text.strip(), f"TXT · {len(text.split())} words"
            except Exception:
                pass
        return None, "Could not decode TXT file"
    elif ext == "docx":
        return extract_docx_text(file_bytes)
    elif ext == "pdf":
        return extract_pdf_text(file_bytes)
    return None, f"Unsupported file type: .{ext}"


# ============================================================
# AI CALL FUNCTIONS
# ============================================================
def call_claude_pagegrid(api_key, model, system, user, progress_cb=None):
    import anthropic
    if not api_key.startswith("sk-pgrid-"):
        raise ValueError("PageGrid key must start with 'sk-pgrid-'.")
    client = anthropic.Anthropic(api_key=api_key, base_url=PAGEGRID_BASE_URL, timeout=300.0)
    full = ""
    with client.messages.stream(
        model=model, max_tokens=16000, system=system,
        messages=[{"role":"user","content":user}],
    ) as stream:
        for chunk in stream.text_stream:
            full += chunk
            if progress_cb:
                progress_cb(len(full), full)
    return full


def call_openai(api_key, model, system, user, progress_cb=None):
    import openai
    client       = openai.OpenAI(api_key=api_key, timeout=300.0)
    is_reasoning = model.startswith("o1") or model.startswith("o3")
    if is_reasoning:
        resp = client.chat.completions.create(
            model=model, max_completion_tokens=16000,
            messages=[{"role":"user","content":f"{system}\n\n---\n\n{user}"}],
        )
        raw = resp.choices[0].message.content
        if progress_cb:
            progress_cb(len(raw), raw)
        return raw
    else:
        full = ""
        for chunk in client.chat.completions.create(
            model=model, max_tokens=16000, stream=True,
            messages=[{"role":"system","content":system},{"role":"user","content":user}],
        ):
            delta = chunk.choices[0].delta.content or ""
            full += delta
            if progress_cb and delta:
                progress_cb(len(full), full)
        return full


def call_gemini(api_key, model, system, user, progress_cb=None):
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    m    = genai.GenerativeModel(model_name=model, system_instruction=system)
    full = ""
    for chunk in m.generate_content(user, stream=True):
        t = getattr(chunk, "text", "") or ""
        full += t
        if progress_cb and t:
            progress_cb(len(full), full)
    return full


def run_generation(api_key, provider, model, system_p, user_p,
                   status_ph, stream_ph=None):
    def _cb(n, text):
        status_ph.markdown(
            f'<div class="stream-progress">'
            f'<b>Generating...</b> &nbsp; {n:,} chars &nbsp;·&nbsp; '
            f'~{n // 5:,} words est. &nbsp;·&nbsp; keep tab open'
            f'</div>',
            unsafe_allow_html=True,
        )
        if stream_ph is not None:
            display = text[-2800:] if len(text) > 2800 else text
            stream_ph.text_area(
                label="stream_live",
                value=display,
                height=270,
                disabled=True,
                label_visibility="collapsed",
            )

    if "PageGrid" in provider:
        return call_claude_pagegrid(api_key, model, system_p, user_p, _cb)
    elif "OpenAI" in provider:
        return call_openai(api_key, model, system_p, user_p, _cb)
    else:
        return call_gemini(api_key, model, system_p, user_p, _cb)


# ============================================================
# THUMBNAIL GENERATION (DALL-E 3)
# ============================================================
def generate_thumbnail_dalle(line1: str, line2: str, line3: str, line4: str,
                              topic: str, api_key: str, variation: int = 0) -> bytes:
    import openai
    client = openai.OpenAI(api_key=api_key, timeout=120.0)

    var_str = f" Creative design variation #{variation}." if variation else ""

    prompt = f"""Professional YouTube thumbnail, SKY Academy Telugu coaching style.{var_str}

STYLE: Photoshop-designed look. NOT AI flashy. Natural lighting. 3-4 colors only. Bold readable fonts. Strong hierarchy. Premium educational brand. 80% content, 20% margins.

TEXT LINES (place exactly as described):

LINE1: "{line1}"
→ Top center. Full-width dark red bar. White bold uppercase. Medium size. Drop shadow.

LINE2: "{line2}"  ← BIGGEST TEXT. MOST DOMINANT.
→ Center of image. Largest font. White or yellow. Yellow/orange strip behind it. Subtle shine. Maximum readability. Visible from 6 feet.

LINE3: "{line3}"
→ Below LINE2. White medium font. No background box. Subtle shadow.

LINE4: "{line4}"
→ Bottom bar. Orange or green filled strip. White bold font inside. High contrast. CTR hook.

BADGE (MANDATORY): తెలుగులో → Bottom-right corner. Yellow rounded badge. Clean shadow.

AUTO THEME (detect topic from lines above, apply matching visuals):
- Polity/Law → Parliament, constitution, court scales
- Geography → India map, globe, terrain
- History → forts, monuments, stone textures
- Economy → rupee, charts, graphs
- Reasoning → arrows, puzzles, number sequences
- Science → atoms, formulas, circuits
- SSC/Exams → exam papers, study desk, target board
- AP High Court → AP High Court building blurred background
- Art/Culture → temples, dancers, sculptures
- Govt Schemes → welfare icons, govt buildings
- Current Affairs → newspaper textures, modern overlays

COLORS (3-4 only, pick by topic):
Dark Blue+Yellow+White | Navy+Orange+White | Teal+Gold+White | Dark Green+Yellow+White
No neon. No rainbow. No excessive glow.

BACKGROUND: Slightly blurred. Cinematic depth. Supports text. No clutter.

LIGHTING: Soft glow only behind LINE2. No bloom. No AI fantasy glow.

SUPPORT VISUALS: Low opacity 15-25%. Subtle. Non-distracting. Auto-match topic.

STRICT RULES:
✗ No faces. No fantasy style. No busy background. No decorative unreadable fonts.
✗ No excessive glow. No clutter. No neon colors.
✓ Strong hierarchy. Clean alignment. Viral-worthy. Mobile readable.

Size: 1792x1024. Ultra sharp. High CTR. SKY Academy premium coaching style."""
response = client.images.generate(model="dall-e-3",
        prompt=prompt,
        size="1792x1024",
        quality="hd",
        response_format="b64_json",
        n=1,
    )
    return base64.b64decode(response.data[0].b64_json)


# ============================================================
# PARSERS
# ============================================================
def parse_single_segment(raw: str):
    cleaned = raw.strip()
    cleaned = re.sub(r"^```[a-zA-Z]*\s*\n?","",cleaned)
    cleaned = re.sub(r"\n?```\s*$","",cleaned).strip()
    for attempt in [cleaned, _sanitize_json_str(cleaned)]:
        try:
            result = json.loads(attempt)
            if isinstance(result, dict):
                return result
            if isinstance(result, list) and result:
                return result[0]
        except Exception:
            pass
    m = re.search(r'\{[\s\S]*\}', cleaned)
    if m:
        for attempt in [m.group(), _sanitize_json_str(m.group())]:
            try:
                result = json.loads(attempt)
                if isinstance(result, dict):
                    return result
            except Exception:
                pass
    return None


def parse_seo_json(raw: str):
    cleaned = raw.strip()
    cleaned = re.sub(r"^```[a-zA-Z]*\s*\n?","",cleaned)
    cleaned = re.sub(r"\n?```\s*$","",cleaned).strip()
    for attempt in [cleaned, _sanitize_json_str(cleaned)]:
        try:
            result = json.loads(attempt)
            if isinstance(result, dict):
                return result
        except Exception:
            pass
    m = re.search(r'\{[\s\S]*\}', cleaned)
    if m:
        for attempt in [m.group(), _sanitize_json_str(m.group())]:
            try:
                return json.loads(attempt)
            except Exception:
                pass
    return None


# ============================================================
# GOOGLE SHEETS
# ============================================================
def push_to_gsheet(chunks, seo_title="", seo_tags=""):
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        scopes = ["https://spreadsheets.google.com/feeds",
                  "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(_HARDCODED_CREDS, scopes=scopes)
        gc    = gspread.authorize(creds)
        sheet = gc.open_by_key(SHEET_ID)
        ws    = sheet.get_worksheet(0)
        ws.batch_clear(["A2:E10000"])
        rows = [[i+1, c.get("title", f"Segment {i+1}"), c.get("telugu_text","")]
                for i, c in enumerate(chunks)]
        ws.update("A2", rows, value_input_option="RAW")
        seo_msg = ""
        if seo_title.strip():
            ws.update("D2", [[seo_title.strip()]], value_input_option="RAW")
            seo_msg += " · Title --> D2"
        if seo_tags.strip():
            ws.update("E2", [[seo_tags.strip()]], value_input_option="RAW")
            seo_msg += " · Tags --> E2"
        return True, f"Pushed {len(rows)} segments (A=Seg, B=Title, C=Script{seo_msg}) -- cleared old data."
    except Exception as exc:
        return False, f"Sheets error: {exc}"


def _handle_api_error(exc, model_choice):
    err = str(exc)
    if "401" in err or "authentication_error" in err:
        st.error("**401 -- Invalid API key.**\n\nPageGrid key must start with `sk-pgrid-`")
    elif "402" in err or "billing_error" in err:
        st.error("**402 -- Wallet balance is $0.**\n\nAdd funds at pagegrid.in")
    elif "404" in err or "not found or inactive" in err:
        st.error(f"**404 -- Model `{model_choice}` not found.**")
    elif "429" in err or "rate_limit" in err:
        st.error("**429 -- Rate limited.** Wait ~60 sec and retry.")
    elif "524" in err or "timeout" in err.lower():
        st.error("**Timeout.** Try a faster model or reduce word count.")
    else:
        st.error(f"Generation error: {exc}")


def _show_manual_recovery(display_topic, display_source):
    st.markdown("#### Manual Recovery")
    manual_raw = st.text_area(
        "Paste Raw JSON here", height=200,
        placeholder='[{"seg":1,"title":"Intro","telugu_text":"..."}]',
        key="manual_json_input",
    )
    if st.button("Retry Parse", key="retry_parse_btn"):
        parsed = parse_segments(manual_raw or st.session_state.raw_response)
        if parsed:
            store_new_chunks(parsed)
            st.session_state.last_topic  = display_topic
            st.session_state.last_source = display_source
            st.success(f"Recovered! {len(parsed)} segments parsed.")
            st.rerun()
        else:
            st.error("Still could not parse. Check for unescaped quotes.")


# ============================================================
# SESSION STATE INIT
# ============================================================
_defaults = {
    "chunks":                 None,
    "raw_response":           "",
    "last_topic":             "",
    "last_source":            "",
    "pdf_text":               "",
    "last_pdf_sig":           "",
    "last_pdf_lib":           "",
    "transcript_extractions": [],
    "transcript_files_sig":   "",
    "seo_pack":               None,
    "handout_data":           None,
    "thumb_img_bytes":        None,
    "thumb_suggestions":      None,
    "thumb_variation":        0,
}
for _k, _v in _defaults.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ============================================================
# SIDEBAR
# ============================================================
with st.sidebar:
    st.markdown("## Configuration")
    st.markdown("### AI Provider & Model")
    provider     = st.selectbox("Provider", list(MODEL_OPTIONS.keys()))
    model_choice = st.selectbox("Model", MODEL_OPTIONS[provider])

    _key_meta = {
        "☁️  Claude  (via PageGrid)": ("PageGrid API Key","sk-pgrid-...","pagegrid.in -> Dashboard -> API Keys","sk-pgrid-"),
        "🟢  OpenAI  (GPT)":          ("OpenAI API Key","sk-...","platform.openai.com -> API Keys","sk-"),
        "🔵  Google  (Gemini)":       ("Google AI API Key","AIzaSy...","aistudio.google.com -> Get API Key","AIzaSy"),
    }
    _lbl, _ph, _hlp, _prefix = _key_meta[provider]
    st.markdown(f"### {_lbl}")
    st.caption(f"Get yours: {_hlp}")

    # Use hardcoded Claude key as default for PageGrid
    _default_key = HARDCODED_CLAUDE_KEY if "PageGrid" in provider else ""
    _api_input = st.text_input(
        _lbl, type="password", placeholder=_ph,
        value=_default_key,
        label_visibility="collapsed",
        key=f"api_key_input_{provider[:4]}",
    )
    # Use hardcoded if field is empty or matches hardcoded
    api_key = _api_input.strip() if _api_input.strip() else _default_key

    if api_key:
        if api_key.startswith(_prefix):
            st.markdown('<p class="key-ok">Key format looks valid</p>', unsafe_allow_html=True)
        else:
            st.markdown(f'<p class="key-warn">Key should start with <code>{_prefix}</code></p>', unsafe_allow_html=True)

    if "PageGrid" in provider:
        st.info("**PageGrid models:**\n- `claude-opus-4-6`\n- `claude-sonnet-4-6`\n- `claude-haiku-4-5`\n\n**base_url:** `https://api.pagegrid.in`", icon="📋")

    st.divider()
    st.markdown("### Google Sheets")
    st.success("Service account loaded\n\n**forscripting@gen-lang-client...**\n\nPush to Sheets always ready.", icon="🔑")
    st.caption(f"[Open Target Sheet](https://docs.google.com/spreadsheets/d/{SHEET_ID}/edit)")

    # ── CHANGE 2: OpenAI key input for DALL-E thumbnail ──
    st.divider()
    st.markdown("### Thumbnail (DALL-E 3)")
    dalle_key_input = st.text_input(
        "OpenAI API Key for DALL-E",
        type="password",
        placeholder="sk-proj-...",
        help="Needs a valid OpenAI key with DALL-E 3 access. Get one at platform.openai.com",
        key="dalle_key_sidebar",
    )
    if dalle_key_input.strip():
        st.markdown('<p class="key-ok">DALL-E key entered ✓</p>', unsafe_allow_html=True)
    else:
        st.markdown('<p class="key-warn">Paste your OpenAI key here to enable thumbnail generation</p>', unsafe_allow_html=True)

    st.divider()
    st.caption("SCRIPT ENGINE v4.0 · SKY Academy Internal Tool")
    st.caption(f"Each segment strictly 150-180 words")


# ============================================================
# LIVE STREAM MASTER PLACEHOLDER
# ============================================================
_stream_master = st.empty()

# ============================================================
# MAIN LAYOUT
# ============================================================
left, right = st.columns([1,1], gap="large")

topic_input      = ""
topic_hint_input = ""
input_mode       = ""
merge_mode_val   = "auto"
video_type       = "subjective"

with left:
    st.markdown("## Script Input")
    st.markdown("### Video Type")
    vtype_label = st.radio(
        "Select video type",
        options=["📚 Subjective  --  Deep Subject Teaching","🎯 General  --  Strategy / Motivation / Guidance"],
        key="video_type_select", label_visibility="collapsed",
    )
    video_type = "general" if vtype_label.startswith("🎯") else "subjective"

    if video_type == "general":
        st.markdown('<div class="vtype-general"><b>General / Strategy Mode</b> -- High motivation, strategy hints, community building.</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="vtype-subjective"><b>Subjective / Deep Teaching Mode</b> -- Max memory hints, exam angles, PYQ references.</div>', unsafe_allow_html=True)

    st.markdown("---")
    input_mode = st.radio(
        "What are you providing?",
        options=["📌 Topic Name  -->  Generate from Scratch",
                 "📝 Competitor Transcripts  -->  Merge to SKY Style",
                 "📚 Book / PDF Section  -->  Convert to SKY Style"],
    )
    st.markdown("---")

    # MODE 1 -- TOPIC
    if input_mode.startswith("📌"):
        st.markdown('<div class="mode-topic"><b>Topic Mode</b> -- Type any subject. SKY Engine writes a complete original voiceover in Telugu lipi.</div>', unsafe_allow_html=True)
        topic_input = st.text_area(
            "Topic / Subject *",
            placeholder="e.g.  Panchayati Raj – 73rd Amendment\n       Photosynthesis Process",
            height=130,
        )
        topic_hint_input = ""

    # MODE 2 -- MULTI-TRANSCRIPT
    elif input_mode.startswith("📝"):
        st.markdown('<div class="mode-transcript"><b>Multi-Transcript Mode</b> -- Upload competitor transcript files. SKY Engine merges and rewrites in Telugu lipi.</div>', unsafe_allow_html=True)
        topic_hint_input = st.text_input("Topic hint  (optional)", placeholder="e.g.  UPSC 2025 Cutoff, TSPSC Exam Date...")
        merge_label = st.radio(
            "How should multiple files be merged?",
            options=["Auto-detect  (AI decides)",
                     "Synthesize same topic  (different data on same subject)",
                     "Merge different topics  (weave different aspects together)"],
        )
        merge_mode_val = {
            "Auto-detect  (AI decides)":                               "auto",
            "Synthesize same topic  (different data on same subject)": "synthesize",
            "Merge different topics  (weave different aspects together)": "merge_aspects",
        }[merge_label]
        transcript_files = st.file_uploader(
            "Upload Transcript Files *", type=["docx","txt","pdf"],
            accept_multiple_files=True,
        )
        topic_input = ""

        if transcript_files:
            new_sig = "|".join(f"{f.name}_{len(f.getvalue())}" for f in transcript_files)
            if new_sig != st.session_state.transcript_files_sig:
                extractions = []
                with st.spinner(f"Extracting text from {len(transcript_files)} file(s)..."):
                    for f in transcript_files:
                        fb = f.getvalue()
                        text, info = extract_any_file(fb, f.name)
                        if text:
                            extractions.append({"filename":f.name,"text":text,"words":len(text.split()),"ok":True,"error":"","info":info})
                        else:
                            extractions.append({"filename":f.name,"text":"","words":0,"ok":False,"error":info,"info":""})
                st.session_state.transcript_extractions = extractions
                st.session_state.transcript_files_sig   = new_sig

            exts      = st.session_state.transcript_extractions
            ok_count  = sum(1 for e in exts if e["ok"])
            err_count = len(exts) - ok_count
            if ok_count:
                total_words = sum(e["words"] for e in exts if e["ok"])
                st.markdown(
                    f'<div class="merge-box"><b>{ok_count} file(s) ready</b>'
                    + (f' · {err_count} failed' if err_count else '')
                    + f' | ~{total_words:,} total words | Merge: <b>{merge_label}</b></div>',
                    unsafe_allow_html=True,
                )
            for e in exts:
                if e["ok"]:
                    st.markdown(f'<div class="fcard-ok">✅ <b>{e["filename"]}</b> -- ~{e["words"]:,} words · {e["info"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="fcard-err">❌ <b>{e["filename"]}</b> -- {e["error"]}</div>', unsafe_allow_html=True)
            ok_exts = [e for e in exts if e["ok"]]
            if ok_exts:
                with st.expander(f"Preview extracted content ({len(ok_exts)} files)", expanded=False):
                    for e in ok_exts:
                        st.markdown(f"**{e['filename']}** -- {e['words']:,} words")
                        st.text(e["text"][:500] + ("\n\n[... more ...]" if len(e["text"])>500 else ""))
                        st.markdown("---")
        elif st.session_state.transcript_extractions:
            st.session_state.transcript_extractions = []
            st.session_state.transcript_files_sig   = ""

    # MODE 3 -- BOOK PDF
    else:
        st.markdown('<div class="mode-pdf"><b>PDF Mode</b> -- Upload a book chapter or study notes PDF. SKY Engine transforms it into Telugu lipi voiceover.</div>', unsafe_allow_html=True)
        topic_hint_input = st.text_input("Topic / Chapter context  (optional)", placeholder="e.g.  Chapter 3: Directive Principles...")
        pdf_file = st.file_uploader("Upload PDF *", type=["pdf"])
        topic_input = ""
        if pdf_file is not None:
            file_bytes = pdf_file.read()
            file_sig   = f"{pdf_file.name}_{len(file_bytes)}"
            if file_sig != st.session_state.last_pdf_sig:
                with st.spinner("Extracting text from PDF..."):
                    extracted, lib_info = extract_pdf_text(file_bytes)
                    if extracted:
                        st.session_state.pdf_text     = extracted
                        st.session_state.last_pdf_sig = file_sig
                        st.session_state.last_pdf_lib = lib_info
                    else:
                        st.session_state.pdf_text     = ""
                        st.session_state.last_pdf_sig = ""
                        st.error(f"Could not extract PDF text. {lib_info}")
            if st.session_state.pdf_text:
                wc = len(st.session_state.pdf_text.split())
                st.success(f"Extracted: ~{wc:,} words · {st.session_state.last_pdf_lib}")
                with st.expander("Preview extracted text", expanded=False):
                    st.text(st.session_state.pdf_text[:800] + "\n\n[... more ...]")
                topic_input = st.session_state.pdf_text

    st.markdown("---")
    approx_words = st.number_input(
        "Approximate Total Script Words",
        min_value=150, max_value=6000, value=660, step=165,
    )
    num_segs = max(1, math.ceil(approx_words / WORDS_PER_SEGMENT))
    est_dur  = round(approx_words / 130, 1)
    st.markdown(
        f'<div class="word-info"><b>{num_segs} segments</b> will be generated &nbsp;·&nbsp; '
        f'~{approx_words} words &nbsp;·&nbsp; ~{est_dur} min video &nbsp;·&nbsp; 150-180w per segment</div>',
        unsafe_allow_html=True,
    )
    special_instructions = st.text_area(
        "Special Instructions  (optional)",
        placeholder="Focus on memory tricks\nTarget: UPSC Mains aspirants",
        height=80,
    )
    c1, c2 = st.columns(2)
    with c1:
        gen_btn   = st.button("Generate Script", type="primary", use_container_width=True)
    with c2:
        clear_btn = st.button("Clear All", use_container_width=True)


# CLEAR
if clear_btn:
    for _k, _v in _defaults.items():
        st.session_state[_k] = _v
    for _i in range(50):
        for _pfx in ["tv_","regen_instr_"]:
            if f"{_pfx}{_i}" in st.session_state:
                del st.session_state[f"{_pfx}{_i}"]
    for _k in ["seo_title_edit","seo_tags_edit","handout_heading_input",
               "handout_youtube_input","handout_pages_input",
               "thumb_l1","thumb_l2","thumb_l3","thumb_l4",
               "quick_notes_heading","quick_notes_yt",
               "standalone_handout_topic"]:
        if _k in st.session_state:
            del st.session_state[_k]
    st.rerun()


# ============================================================
# GENERATE
# ============================================================
if gen_btn:
    if   input_mode.startswith("📌"): mode_name = "topic"
    elif input_mode.startswith("📝"): mode_name = "transcript"
    else:                              mode_name = "pdf"

    if not api_key.strip():
        _stream_master.empty()
        st.error("Please enter your API key in the sidebar!")

    elif mode_name == "topic" and not topic_input.strip():
        _stream_master.empty()
        st.error("Please enter a topic!")

    elif mode_name == "transcript" and not [e for e in st.session_state.transcript_extractions if e["ok"]]:
        _stream_master.empty()
        st.error("Please upload at least one transcript file (.docx/.txt/.pdf).")

    elif mode_name == "pdf" and not topic_input.strip():
        _stream_master.empty()
        st.error("Please upload a PDF. Make sure text was extracted successfully.")

    else:
        with _stream_master.container():
            st.markdown("""
<div class="live-stream-box">
  <div class="live-stream-title">Live Generation Stream</div>
  <div class="live-stream-hint">
    Watching the AI write your script in real-time. All Telugu output will be in Telugu Unicode script (lipi).
    150-180 words per segment. The preview panel updates automatically when generation completes.
  </div>
</div>
""", unsafe_allow_html=True)
            _live_stat = st.empty()
            _live_txt  = st.empty()

        vtype_disp = "General/Strategy" if video_type=="general" else "Subjective/Teaching"

        try:
            if mode_name == "transcript":
                ok_exts          = [e for e in st.session_state.transcript_extractions if e["ok"]]
                safe_transcripts = [{"filename":e["filename"],"text":e["text"]} for e in ok_exts]
                st.info(f"Merging {len(safe_transcripts)} transcript(s) → SKY Academy Telugu lipi [{vtype_disp}]... keep tab open", icon="⏳")
                system_p, user_p = build_prompts_multi_transcript(
                    safe_transcripts, topic_hint_input,
                    num_segs, special_instructions, merge_mode_val, video_type,
                )
                display_topic  = topic_hint_input.strip()[:60] or "Merged Script"
                display_source = "📝 Transcript --> SKY"

            elif mode_name == "topic":
                st.info(f"Generating {num_segs} segments [{vtype_disp}] via {model_choice}... keep tab open", icon="⏳")
                system_p, user_p = build_prompts_topic(
                    topic_input, num_segs, special_instructions, video_type
                )
                display_topic  = topic_input.strip()[:60]
                display_source = "📌 Topic"

            else:
                st.info(f"Converting PDF content → SKY Academy Telugu lipi [{vtype_disp}]... keep tab open", icon="⏳")
                system_p, user_p = build_prompts_pdf(
                    topic_input, topic_hint_input, num_segs, special_instructions, video_type
                )
                display_topic  = topic_hint_input.strip()[:60] or "PDF Content"
                display_source = "📚 PDF --> SKY"

            raw = run_generation(
                api_key, provider, model_choice,
                system_p, user_p,
                _live_stat, _live_txt,
            )
            _stream_master.empty()
            st.session_state.raw_response = raw
            parsed = parse_segments(raw)

            if parsed:
                if len(parsed) < num_segs:
                    st.warning(
                        f"Partial recovery: got {len(parsed)} of {num_segs} segments. "
                        f"Try reducing word count or using a faster model.",
                        icon="⚠️"
                    )
                store_new_chunks(parsed)
                st.session_state.last_topic  = display_topic
                st.session_state.last_source = display_source
                st.success(f"{len(parsed)} segments generated! [{display_source}] · Telugu lipi output")
                st.rerun()
            else:
                st.error("Could not parse JSON from AI response. Try the manual recovery below.")
                _show_manual_recovery(display_topic, display_source)

        except Exception as exc:
            _stream_master.empty()
            _handle_api_error(exc, model_choice)


# ============================================================
# RIGHT: PREVIEW
# ============================================================
with right:
    st.markdown("## Preview")
    chunks = st.session_state.chunks

    if chunks:
        total_words = sum(len(c.get("telugu_text","").split()) for c in chunks)
        est_min     = round(total_words / 130, 1)
        badge_class = {
            "📌 Topic":               "badge-topic",
            "📝 Transcript --> SKY":  "badge-transcript",
            "📚 PDF --> SKY":         "badge-pdf",
        }.get(st.session_state.last_source, "badge-topic")
        vtype_badge_color = "#f97316" if video_type=="general" else "#22c55e"
        vtype_badge_label = "General" if video_type=="general" else "Subjective"

        st.markdown(
            f'<span class="{badge_class}">{st.session_state.last_source}</span>'
            f'&nbsp;&nbsp;'
            f'<span style="background:{vtype_badge_color};color:#fff;border-radius:8px;padding:4px 12px;font-size:.8rem;font-weight:700">{vtype_badge_label}</span>'
            f'&nbsp;&nbsp;'
            f'<span class="stat-pill">{len(chunks)} segments</span>'
            f'<span class="stat-pill">~{total_words:,} words</span>'
            f'<span class="stat-pill">~{est_min} min</span>'
            f'<span class="stat-pill" style="background:linear-gradient(135deg,#dcfce7,#bbf7d0);color:#14532d;">Telugu Lipi</span>',
            unsafe_allow_html=True,
        )
        st.markdown("")

        tabs = st.tabs([f"▶ {i+1}" for i in range(len(chunks))])
        for tab, chunk, idx in zip(tabs, chunks, range(len(chunks))):
            with tab:
                seg_words = len(chunk.get("telugu_text","").split())
                word_color = "#16a34a" if 150 <= seg_words <= 180 else "#f97316"
                st.caption(
                    f"**{chunk.get('title', f'Segment {idx+1}')}** &nbsp;·&nbsp; "
                    f"<span style='color:{word_color};font-weight:700'>{seg_words} words</span>"
                    + (" · Hook + Welcome" if idx==0 else "")
                    + (" · Closing CTA" if idx==len(chunks)-1 else ""),
                    unsafe_allow_html=True,
                )
                st.markdown("**Voiceover Script (Telugu Lipi)**")
                st.text_area(f"vo_{idx}", key=f"tv_{idx}", height=280, label_visibility="collapsed")

                with st.expander(f"🔄 Redo Segment {idx+1}", expanded=False):
                    st.markdown('<div class="regen-hint">Describe what to change -- only this segment will be regenerated. Output stays in Telugu lipi.</div>', unsafe_allow_html=True)
                    regen_instr = st.text_area(
                        "Change instruction",
                        placeholder="e.g. Add 2 memory hints\nMake more energetic\nToo long -- trim",
                        height=90, key=f"regen_instr_{idx}", label_visibility="collapsed",
                    )
                    regen_btn = st.button(f"Regenerate Segment {idx+1}", key=f"regen_btn_{idx}", use_container_width=True)
                    if regen_btn:
                        if not api_key.strip():
                            st.error("Enter API key in sidebar first!")
                        else:
                            _rs  = st.empty()
                            _rst = st.empty()
                            with st.spinner(f"Regenerating segment {idx+1}..."):
                                try:
                                    _instr = st.session_state.get(f"regen_instr_{idx}","") or "Improve -- better flow, sharper memory hints"
                                    sys_r, usr_r = build_regen_segment_prompt(
                                        video_type, st.session_state.last_topic,
                                        st.session_state.chunks, idx, _instr,
                                        len(st.session_state.chunks),
                                    )
                                    raw_r = run_generation(api_key, provider, model_choice, sys_r, usr_r, _rs, _rst)
                                    _rs.empty()
                                    _rst.empty()
                                    new_chunk = parse_single_segment(raw_r)
                                    if new_chunk and isinstance(new_chunk, dict):
                                        cleaned_tv = strip_emojis(new_chunk.get("telugu_text",""))
                                        new_chunk["telugu_text"] = cleaned_tv
                                        st.session_state.chunks[idx] = new_chunk
                                        st.session_state[f"tv_{idx}"] = cleaned_tv
                                        st.success(f"Segment {idx+1} regenerated!")
                                        st.rerun()
                                    else:
                                        st.error("Could not parse regenerated segment. Try again.")
                                except Exception as _exc:
                                    _rs.empty()
                                    _rst.empty()
                                    st.error(f"Regen error: {_exc}")

        with st.expander("Full Script -- Continuous Flow", expanded=False):
            full = "\n\n".join(
                f"--- [{chunk.get('title', f'Segment {i+1}')}] ---\n"
                + st.session_state.get(f"tv_{i}", chunk.get("telugu_text",""))
                for i, chunk in enumerate(chunks)
            )
            st.text_area("full_script", value=full, height=400, label_visibility="collapsed")

        st.divider()
        st.markdown("#### YouTube SEO Pack")
        st.markdown('<div class="seo-box">Generate an optimized YouTube title and 20 tags. Title → <b>D2</b>, Tags → <b>E2</b> when pushed to Sheets.</div>', unsafe_allow_html=True)
        seo_gen_btn = st.button("Generate SEO Pack", key="seo_gen_btn", use_container_width=True)
        if seo_gen_btn:
            if not api_key.strip():
                st.error("Enter API key in sidebar first!")
            else:
                _ss  = st.empty()
                with st.spinner("Generating YouTube SEO Pack..."):
                    try:
                        sys_seo, usr_seo = build_seo_prompt(st.session_state.last_topic, sync_edits_to_chunks(), video_type)
                        raw_seo = run_generation(api_key, provider, model_choice, sys_seo, usr_seo, _ss)
                        _ss.empty()
                        seo_data = parse_seo_json(raw_seo)
                        if seo_data and isinstance(seo_data, dict):
                            st.session_state.seo_pack = seo_data
                            st.session_state["seo_title_edit"] = seo_data.get("title","")
                            st.session_state["seo_tags_edit"]  = seo_data.get("tags","")
                            st.success("SEO Pack generated!")
                            st.rerun()
                        else:
                            st.error("Could not parse SEO response. Try again.")
                    except Exception as _exc:
                        _ss.empty()
                        st.error(f"SEO error: {_exc}")

        if st.session_state.seo_pack:
            with st.expander("Edit SEO Pack before pushing to Sheets", expanded=True):
                st.text_input("YouTube Title  -->  D2", key="seo_title_edit")
                st.text_area("YouTube Tags  -->  E2  (comma-separated)", key="seo_tags_edit", height=80)
                st.caption(f"Title: {len(st.session_state.get('seo_title_edit',''))}/70 chars")

        st.divider()

        st.markdown("#### Quick Study Notes Download")
        qn_heading = st.text_input(
            "PDF Heading (optional)",
            placeholder="e.g.  Preamble of Indian Constitution",
            key="quick_notes_heading",
        )
        qn_ytlink = st.text_input(
            "YouTube Video Link (optional)",
            placeholder="https://youtu.be/...",
            key="quick_notes_yt",
        )

        _fname      = st.session_state.last_topic[:20].replace(" ","_")
        synced_data = sync_edits_to_chunks()

        ba, bb, bc = st.columns(3)
        with ba:
            st.download_button(
                "Download JSON",
                data=json.dumps(synced_data, ensure_ascii=False, indent=2).encode("utf-8"),
                file_name=f"sky_script_{_fname}.json",
                mime="application/json",
                use_container_width=True,
            )
        with bb:
            txt_out = "\n\n".join(
                f"--- [{c.get('title',f'Segment {i+1}')}] ---\n{c.get('telugu_text','')}"
                for i, c in enumerate(synced_data)
            )
            st.download_button(
                "Download TXT",
                data=txt_out.encode("utf-8"),
                file_name=f"sky_script_{_fname}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with bc:
            notes_html = generate_study_notes_html(
                synced_data, video_type,
                heading=st.session_state.get("quick_notes_heading",""),
                youtube_link=st.session_state.get("quick_notes_yt",""),
            )
            st.download_button(
                "Download Study Notes",
                data=notes_html.encode("utf-8"),
                file_name=f"sky_notes_{_fname}.html",
                mime="text/html",
                use_container_width=True,
                help="Open in Chrome/Edge --> Ctrl+P --> Save as PDF",
            )

        st.caption("Open downloaded .html in Chrome → Ctrl+P → Save as PDF → share on Telegram!")

        push_btn = st.button("Push to Sheets", use_container_width=True, type="primary")
        if push_btn:
            seo_title = st.session_state.get("seo_title_edit","")
            seo_tags  = st.session_state.get("seo_tags_edit","")
            with st.spinner("Clearing old data and writing to Sheet1..."):
                ok, msg = push_to_gsheet(synced_data, seo_title, seo_tags)
            if ok:
                st.success(msg)
                st.markdown(f"[Open Sheet](https://docs.google.com/spreadsheets/d/{SHEET_ID}/edit)")
            else:
                st.error(msg)

    else:
        st.markdown("""
        <div class="empty-preview">
            <h3>Preview will appear here</h3>
            <p style="margin-top:12px;font-size:0.9rem;">
                Choose video type → Input mode → Provide content → Click <b>Generate Script</b><br>
                <span style="font-size:0.8rem;color:#818cf8;">Output: Telugu Unicode lipi · 150-180 words per segment</span>
            </p>
        </div>""", unsafe_allow_html=True)


# ============================================================
# THUMBNAIL GENERATOR -- FULL WIDTH AMBER BLOCK
# ============================================================
st.markdown("---")
st.markdown("""
<div class="thumb-block">
  <div class="thumb-title">🎨 THUMBNAIL GENERATOR</div>
</div>
""", unsafe_allow_html=True)

with st.container():
    st.markdown("""
    <div class="thumb-hint">
    Create high-CTR SKY Academy style thumbnails with DALL-E 3.
    Fill the 4 lines manually, or click <b>AI Suggest Lines</b> to auto-fill based on your script.
    Then click <b>Generate Thumbnail</b> to create a professional Photoshop-style thumbnail.
    <br><b>Note:</b> Paste your OpenAI API key in the sidebar (Thumbnail section) before generating.
    </div>
    """, unsafe_allow_html=True)

    # Show current AI suggestions if available
    if st.session_state.thumb_suggestions:
        sugg = st.session_state.thumb_suggestions
        st.markdown(
            '<div class="thumb-sugg-box">'
            f'<b>AI Suggestions applied:</b> &nbsp; '
            f'L1: <b>{sugg.get("line1","")}</b> &nbsp;|&nbsp; '
            f'L2: <b>{sugg.get("line2","")}</b> &nbsp;|&nbsp; '
            f'L3: <b>{sugg.get("line3","")}</b> &nbsp;|&nbsp; '
            f'L4: <b>{sugg.get("line4","")}</b>'
            '</div>',
            unsafe_allow_html=True,
        )

    # 4 text inputs in 2 columns
    tl_col1, tl_col2 = st.columns(2)
    with tl_col1:
        thumb_l1 = st.text_input(
            "📌 Line 1 — Target Audience / Exam Name",
            placeholder="AP HIGH COURT 2026",
            key="thumb_l1",
        )
        thumb_l2 = st.text_input(
            "🎯 Line 2 — Core Video Topic",
            placeholder="BNS COMPLETE SERIES",
            key="thumb_l2",
        )
    with tl_col2:
        thumb_l3 = st.text_input(
            "📝 Line 3 — Brief Explanation",
            placeholder="ALL TOPICS COVERED | WITH PYQS",
            key="thumb_l3",
        )
        thumb_l4 = st.text_input(
            "⚡ Line 4 — Telugu Punch Line (in Telugu script)",
            placeholder="Miss అవ్వాండి!",
            key="thumb_l4",
        )

    st.caption(
        "**Line 1**: Exam name | **Line 2**: Topic | **Line 3**: Sub-info | **Line 4**: Telugu punch line · "
        "Click 'AI Suggest' to auto-fill from your script"
    )

    tb1, tb2, tb3 = st.columns(3)
    with tb1:
        thumb_suggest_btn = st.button(
            "💡 AI Suggest Lines",
            key="thumb_suggest_btn",
            use_container_width=True,
            help="Uses your generated script to suggest high-CTR text for each line",
        )
    with tb2:
        thumb_gen_btn = st.button(
            "🎨 Generate Thumbnail",
            key="thumb_gen_btn",
            use_container_width=True,
            help="Creates thumbnail with DALL-E 3 (requires OpenAI key in sidebar)",
        )
    with tb3:
        thumb_regen_btn = st.button(
            "🔄 Regenerate (New Style)",
            key="thumb_regen_btn",
            use_container_width=True,
            disabled=st.session_state.thumb_img_bytes is None,
            help="Generate a different design variation",
        )

    # Handle Suggest Lines
    if thumb_suggest_btn:
        if not api_key.strip():
            st.error("Enter API key in sidebar first!")
        else:
            _ts = st.empty()
            with st.spinner("Getting AI line suggestions..."):
                try:
                    sys_ts, usr_ts = build_thumbnail_suggest_prompt(
                        st.session_state.last_topic or "Competitive Exam Preparation",
                        st.session_state.chunks or [],
                        video_type,
                    )
                    raw_ts = run_generation(api_key, provider, model_choice, sys_ts, usr_ts, _ts)
                    _ts.empty()
                    sugg_data = parse_seo_json(raw_ts)
                    if sugg_data and isinstance(sugg_data, dict):
                        st.session_state.thumb_suggestions = sugg_data
                        st.session_state["thumb_l1"] = sugg_data.get("line1", "")
                        st.session_state["thumb_l2"] = sugg_data.get("line2", "")
                        st.session_state["thumb_l3"] = sugg_data.get("line3", "")
                        st.session_state["thumb_l4"] = sugg_data.get("line4", "")
                        st.success("Lines suggested and filled! Edit if needed, then Generate.")
                        st.rerun()
                    else:
                        st.error("Could not parse suggestions. Try again.")
                except Exception as _exc:
                    _ts.empty()
                    st.error(f"Suggestion error: {_exc}")

    # ── CHANGE 3a: Handle Generate Thumbnail (uses dalle_key_sidebar) ──
    if thumb_gen_btn:
        l1 = st.session_state.get("thumb_l1", "").strip()
        l2 = st.session_state.get("thumb_l2", "").strip()
        l3 = st.session_state.get("thumb_l3", "").strip()
        l4 = st.session_state.get("thumb_l4", "").strip()
        dalle_api_key = st.session_state.get("dalle_key_sidebar", "").strip()
        if not dalle_api_key:
            st.error("Please enter your OpenAI API key in the sidebar (under **Thumbnail** section) before generating!")
        elif not any([l1, l2, l3, l4]):
            st.error("Please fill at least one line before generating!")
        else:
            with st.spinner("Generating thumbnail with DALL-E 3 HD... (20-30 seconds)"):
                try:
                    img_bytes = generate_thumbnail_dalle(
                        l1 or "SKY ACADEMY",
                        l2 or "COMPETITIVE EXAM PREPARATION",
                        l3 or "STRATEGY + PDF FREE",
                        l4 or "ఇప్పుడే చూడండి!",
                        st.session_state.last_topic or "competitive exam",
                        api_key=dalle_api_key,
                        variation=0,
                    )
                    st.session_state.thumb_img_bytes = img_bytes
                    st.session_state.thumb_variation = 0
                    st.success("Thumbnail generated!")
                    st.rerun()
                except Exception as _exc:
                    st.error(f"Thumbnail generation error: {_exc}")
                    st.caption("Check that your OpenAI key is valid and has DALL-E 3 access.")

    # ── CHANGE 3b: Handle Regenerate (uses dalle_key_sidebar) ──
    if thumb_regen_btn and st.session_state.thumb_img_bytes is not None:
        l1 = st.session_state.get("thumb_l1", "").strip()
        l2 = st.session_state.get("thumb_l2", "").strip()
        l3 = st.session_state.get("thumb_l3", "").strip()
        l4 = st.session_state.get("thumb_l4", "").strip()
        dalle_api_key = st.session_state.get("dalle_key_sidebar", "").strip()
        new_variation = (st.session_state.thumb_variation + 1) % 50 + 1
        if not dalle_api_key:
            st.error("Please enter your OpenAI API key in the sidebar (under **Thumbnail** section)!")
        else:
            with st.spinner(f"Regenerating thumbnail (variation #{new_variation})..."):
                try:
                    img_bytes = generate_thumbnail_dalle(
                        l1 or "SKY ACADEMY",
                        l2 or "COMPETITIVE EXAM PREPARATION",
                        l3 or "STRATEGY + PDF FREE",
                        l4 or "ఇప్పుడే చూడండి!",
                        st.session_state.last_topic or "competitive exam",
                        api_key=dalle_api_key,
                        variation=new_variation,
                    )
                    st.session_state.thumb_img_bytes = img_bytes
                    st.session_state.thumb_variation = new_variation
                    st.success(f"New variation #{new_variation} generated!")
                    st.rerun()
                except Exception as _exc:
                    st.error(f"Regen error: {_exc}")

    # Display generated thumbnail
    if st.session_state.thumb_img_bytes:
        st.markdown("##### Generated Thumbnail Preview")
        st.image(st.session_state.thumb_img_bytes, use_container_width=True)

        td1, td2 = st.columns(2)
        with td1:
            _tfname = (st.session_state.last_topic[:20] or "thumbnail").replace(" ","_")
            st.download_button(
                "⬇️ Download Thumbnail (PNG)",
                data=st.session_state.thumb_img_bytes,
                file_name=f"sky_thumbnail_{_tfname}.png",
                mime="image/png",
                use_container_width=True,
            )
        with td2:
            st.caption(
                "💡 **Tips:** Download and open in Canva or Photoshop to add exact Telugu text "
                "overlays and fine-tune positioning. DALL-E sets the background style, "
                "you perfect the text."
            )


# ============================================================
# AI HANDOUT GENERATOR -- FULL WIDTH PURPLE BLOCK
# ============================================================
st.markdown("---")
st.markdown("""
<div class="handout-block">
  <div class="handout-title">📄 AI HANDOUT GENERATOR</div>
</div>
""", unsafe_allow_html=True)

with st.container():
    st.markdown("""
    <div class="handout-hint">
    Generate a rich, print-ready classroom handout. Choose <b>From Generated Script</b> to use
    your current session script, or <b>Standalone Mode</b> to generate a handout from any topic/directions
    without needing to generate a script first.
    </div>
    """, unsafe_allow_html=True)

    # Source selector
    handout_source_radio = st.radio(
        "Handout Content Source:",
        options=[
            "📄 From Generated Script  (uses current session)",
            "✍️ Standalone Mode  —  My Own Topic / Directions",
        ],
        key="handout_source_radio",
        horizontal=True,
    )
    use_script_for_handout = handout_source_radio.startswith("📄")

    standalone_topic_val = ""
    if not use_script_for_handout:
        st.markdown("**Describe your topic and what to include:**")
        standalone_topic_val = st.text_area(
            "Topic Overview & Directions",
            placeholder=(
                "Example:\n"
                "Topic: Indian Classical Dance Forms\n"
                "Include: All 8 classical dances recognized by Sangeet Natak Akademi\n"
                "Details for each: state, founder, key features, famous performers\n"
                "Exam focus: UPSC/APPSC/SSC questions on classical dances\n"
                "Language: English only\n"
                "Pages: 2 A4 pages"
            ),
            height=160,
            key="standalone_handout_topic",
        )
    else:
        if not st.session_state.chunks:
            st.warning(
                "No script generated yet. Generate a script first, or switch to **Standalone Mode** "
                "to create a handout without a script.",
                icon="⚠️",
            )

    hcol1, hcol2, hcol3 = st.columns([2, 2, 1])
    with hcol1:
        handout_heading = st.text_input(
            "📝 Handout Heading",
            placeholder="e.g.  Preamble of Indian Constitution",
            key="handout_heading_input",
        )
    with hcol2:
        handout_youtube = st.text_input(
            "🎬 YouTube Video Link (optional)",
            placeholder="https://youtu.be/xxxx  OR  leave blank",
            key="handout_youtube_input",
        )
    with hcol3:
        handout_pages = st.number_input(
            "📄 Pages",
            min_value=1, max_value=8, value=2, step=1,
            key="handout_pages_input",
            help="Approximate number of A4 pages",
        )

    hb1, hb2, hb3 = st.columns(3)
    with hb1:
        gen_handout_btn = st.button(
            "✨ Generate AI Handout",
            use_container_width=True,
            key="gen_handout_btn",
        )
    with hb2:
        if st.session_state.handout_data:
            _hfname = st.session_state.last_topic[:20].replace(" ","_") or "handout"
            _hhtml  = render_handout_html(
                st.session_state.get("handout_heading_input",""),
                st.session_state.get("handout_youtube_input",""),
                st.session_state.handout_data,
                st.session_state.last_topic or "SKY Academy Notes",
            )
            st.download_button(
                "📥 Download PDF (HTML)",
                data=_hhtml.encode("utf-8"),
                file_name=f"sky_handout_{_hfname}.html",
                mime="text/html",
                use_container_width=True,
                help="Open in Chrome → Ctrl+P → Save as PDF",
                key="dl_handout_html",
            )
        else:
            st.button("📥 Download PDF (HTML)", disabled=True, use_container_width=True, key="dl_handout_html_dis")

    with hb3:
        if st.session_state.handout_data:
            _hfname  = st.session_state.last_topic[:20].replace(" ","_") or "handout"
            _docx_bytes = render_handout_docx(
                st.session_state.get("handout_heading_input",""),
                st.session_state.get("handout_youtube_input",""),
                st.session_state.handout_data,
                st.session_state.last_topic or "SKY Academy Notes",
            )
            if _docx_bytes:
                st.download_button(
                    "📄 Download DOCX",
                    data=_docx_bytes,
                    file_name=f"sky_handout_{_hfname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_handout_docx",
                )
            else:
                st.button("📄 Download DOCX", disabled=True, use_container_width=True,
                          help="Install python-docx: pip install python-docx",
                          key="dl_handout_docx_dis")
        else:
            st.button("📄 Download DOCX", disabled=True, use_container_width=True, key="dl_handout_docx_dis2")

    # GENERATE HANDOUT
    if gen_handout_btn:
        if not api_key.strip():
            st.error("Enter your API key in the sidebar first!")
        elif use_script_for_handout and not st.session_state.chunks:
            st.error("No script generated yet! Generate a script first, or switch to Standalone Mode.")
        elif not use_script_for_handout and not standalone_topic_val.strip():
            st.error("Please provide your topic overview and directions!")
        else:
            _hs  = st.empty()
            _hst = st.empty()
            with st.spinner(f"Generating {handout_pages}-page AI handout..."):
                try:
                    if use_script_for_handout:
                        synced_for_handout = sync_edits_to_chunks()
                        _htopic = st.session_state.last_topic or "Competitive Exam Topic"
                        sys_h, usr_h = build_handout_prompt(
                            _htopic,
                            synced_for_handout,
                            handout_pages,
                        )
                    else:
                        # Standalone mode -- use topic overview
                        _htopic = st.session_state.get("standalone_handout_topic","").strip()[:60] or "Study Topic"
                        sys_h, usr_h = build_standalone_handout_prompt(
                            standalone_topic_val,
                            handout_pages,
                        )
                        # Update last_topic for download filename
                        if not st.session_state.last_topic:
                            st.session_state.last_topic = _htopic

                    raw_h = run_generation(api_key, provider, model_choice, sys_h, usr_h, _hs, _hst)
                    _hs.empty()
                    _hst.empty()
                    handout_content = parse_handout_content(raw_h)
                    if handout_content and isinstance(handout_content, dict):
                        st.session_state.handout_data = handout_content
                        st.success(
                            f"Handout generated! {len(handout_content.get('sections',[]))} sections · "
                            f"~{handout_pages} pages. Use the Download buttons above."
                        )
                        st.rerun()
                    else:
                        st.error("Could not parse handout content. Try again or use a different model.")
                        with st.expander("Raw handout response (debug)", expanded=False):
                            st.code(raw_h[:3000], language="json")
                except Exception as _exc:
                    _hs.empty()
                    _hst.empty()
                    _handle_api_error(_exc, model_choice)

    if st.session_state.handout_data:
        with st.expander("📋 Preview Handout Structure", expanded=False):
            hd = st.session_state.handout_data
            st.markdown(f"**Topic Title:** {hd.get('topic_title','')}")
            st.markdown(f"**Exam Importance:** {hd.get('exam_importance','')[:200]}...")
            st.markdown(f"**Key Facts:** {len(hd.get('key_facts',[]))} items")
            st.markdown(f"**Sections:** {len(hd.get('sections',[]))}")
            for s in hd.get("sections",[]):
                has_table = "✅ Table" if s.get("table") else ""
                has_qs    = f"✅ {len(s.get('questions',[]))} PYQs" if s.get("questions") else ""
                st.markdown(
                    f"- **{s.get('heading','')}** · "
                    f"{len(s.get('bullets',[]))} bullets {has_table} {has_qs}"
                )

    st.caption("💡 Open downloaded HTML in Chrome or Edge → Ctrl+P → Save as PDF. Footer appears on every printed page.")


# ============================================================
# RAW RESPONSE DEBUG
# ============================================================
if st.session_state.raw_response:
    with st.expander("Raw AI Response  (debug / manual copy-paste fallback)", expanded=False):
        st.code(st.session_state.raw_response[:8000], language="json")


# ============================================================
# FOOTER
# ============================================================
st.divider()
st.markdown(
    "<div style='text-align:center;color:#aaa;font-size:11px;'>"
    "<span style='color:#FF6B6B;font-weight:800;'>S</span>"
    "<span style='color:#FFE66D;font-weight:800;'>K</span>"
    "<span style='color:#4ECDC4;font-weight:800;'>Y</span>"
    " <b>ACADEMY</b> &nbsp;|&nbsp; SCRIPT ENGINE v4.0 &nbsp;|&nbsp; "
    "Telugu Lipi · 150-180w Chunks · DALL-E 3 Thumbnails · Standalone Handout &nbsp;|&nbsp; "
    "Powered by PageGrid + Anthropic SDK"
    "</div>",
    unsafe_allow_html=True,
)
